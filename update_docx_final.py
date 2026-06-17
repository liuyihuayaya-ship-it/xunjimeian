#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Update the 常用运算符和数学函数 docx with ch9 I/O content."""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree
import copy

docx_path = r'C:\Users\lenovo\Desktop\常用运算符和数学函数_含示例解释.docx'
doc = Document(docx_path)

# Reference formatting from existing document
# Section headers: bold, size=139700 (11pt), with specific font
# Subsection headers: bold, size=139700
# Table header rows: bold text
# Code/comments: size=127000 (10pt)

def make_bold_paragraph(doc, text, font_size=Pt(11)):
    """Create a bold paragraph matching existing section style."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = font_size
    run.font.name = '等线'
    rPr = run._r.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="等线"/>')
    rPr.insert(0, rFonts)
    return p

def make_normal_paragraph(doc, text, font_size=Pt(10)):
    """Create a normal paragraph matching existing style."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = font_size
    run.font.name = '等线'
    rPr = run._r.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="等线"/>')
    rPr.insert(0, rFonts)
    return p

def create_table(doc, headers, rows):
    """Create a formatted table matching existing table style."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ''
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = '等线'

    # Data rows
    for r, row_data in enumerate(rows):
        row_cells = table.rows[r + 1].cells
        for c, cell_text in enumerate(row_data):
            row_cells[c].text = ''
            p = row_cells[c].paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(10)
            run.font.name = '等线'

    return table

# === Now find the insertion point ===
# We need to find P43 (Section 5) element in the document body XML
# and insert before it

body = doc.element.body

# Find all paragraph elements
nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

# Find the paragraph that contains "五、" (Section 5)
p43_element = None
for p in body.findall('.//' + qn('w:p')):
    texts = p.findall('.//' + qn('w:t'))
    combined = ''.join(t.text or '' for t in texts)
    if '五、常用数学函数' in combined:
        p43_element = p
        break

if p43_element is None:
    print("ERROR: Could not find Section 5 paragraph!")
    sys.exit(1)

print(f"Found Section 5 element at index: {list(body).index(p43_element)}")

# Collect all elements to insert (they'll be added before P43)
elements_to_insert = []

# === 4.3 I/O 操纵符 ===
# Create paragraph for 4.3
p43_header = make_bold_paragraph(doc, '4.3 I/O 操纵符（需 #include <iomanip>）')
elements_to_insert.append(p43_header)

# Create empty paragraph for spacing
sp1 = make_normal_paragraph(doc, '')
elements_to_insert.append(sp1)

# Table for I/O manipulators
manip_headers = ['操纵符', '功能', '示例']
manip_rows = [
    ['setw(n)', '设置输出宽度为 n 个字符', 'cout << setw(10) << x;'],
    ['setfill(c)', '设置填充字符为 c', 'cout << setfill(\'*\') << setw(8) << x;'],
    ['setprecision(n)', '设置浮点数精度为 n 位', 'cout << setprecision(6) << pi;'],
    ['left', '输出左对齐', 'cout << left << setw(10) << x;'],
    ['right', '输出右对齐（默认）', 'cout << right << setw(10) << x;'],
    ['fixed', '以固定小数点方式输出', 'cout << fixed << 3.14159; // 3.141590'],
    ['scientific', '以科学计数法输出', 'cout << scientific << 3.14; // 3.140000e+000'],
    ['hex', '以十六进制输出整数', 'cout << hex << 255; // ff'],
    ['oct', '以八进制输出整数', 'cout << oct << 255; // 377'],
    ['dec', '以十进制输出整数（默认）', 'cout << dec << 255; // 255'],
    ['showbase', '显示进制前缀', 'cout << showbase << hex << 255; // 0xff'],
    ['showpoint', '始终显示小数点', 'cout << showpoint << 5.0; // 5.00000'],
    ['endl', '换行并刷新输出缓冲区', 'cout << "Hello" << endl;'],
    ['\\n', '仅换行，不刷新（更快）', 'cout << "Hello\\n";'],
]
manip_table = create_table(doc, manip_headers, manip_rows)
elements_to_insert.append(manip_table)

# Code example for manipulators
manip_example = make_normal_paragraph(doc, '')
elements_to_insert.append(manip_example)

manip_code1 = make_normal_paragraph(doc, '// 格式控制示例：')
elements_to_insert.append(manip_code1)

manip_code2 = make_normal_paragraph(doc, 'int inum = 255;')
elements_to_insert.append(manip_code2)

manip_code3 = make_normal_paragraph(doc, 'cout << hex << showbase << inum;       // 输出: 0xff(十六进制带前缀)')
elements_to_insert.append(manip_code3)

manip_code4 = make_normal_paragraph(doc, 'cout << right << setw(10) << setfill(\'*\') << inum;  // 输出: *******255')
elements_to_insert.append(manip_code4)

manip_code5 = make_normal_paragraph(doc, 'double pi = 3.1415926535;')
elements_to_insert.append(manip_code5)

manip_code6 = make_normal_paragraph(doc, 'cout << fixed << setprecision(6) << pi;  // 输出: 3.141593（固定6位小数）')
elements_to_insert.append(manip_code6)

manip_code7 = make_normal_paragraph(doc, 'cout << scientific << setprecision(9) << pi;  // 输出: 3.141592654e+000')
elements_to_insert.append(manip_code7)

# === 4.4 格式控制标志 ===
elements_to_insert.append(make_normal_paragraph(doc, ''))
p44_header = make_bold_paragraph(doc, '4.4 格式控制标志（ios::flags）')
elements_to_insert.append(p44_header)
elements_to_insert.append(make_normal_paragraph(doc, ''))

# Table for format flags
flags_headers = ['标志', '功能', '说明']
flags_rows = [
    ['ios::skipws', '跳过空白字符', '输入时跳过前导空格/制表符（默认启用）'],
    ['ios::left', '左对齐输出', '在设定宽度内靠左显示'],
    ['ios::right', '右对齐输出', '在设定宽度内靠右显示（默认）'],
    ['ios::internal', '内部对齐', '符号靠左，数值靠右，中间填充'],
    ['ios::dec', '十进制', '整数以十进制显示（默认）'],
    ['ios::oct', '八进制', '整数以八进制显示'],
    ['ios::hex', '十六进制', '整数以十六进制显示'],
    ['ios::showbase', '显示进制前缀', '八进制加0，十六进制加0x'],
    ['ios::showpoint', '显示小数点', '浮点数始终显示小数点'],
    ['ios::uppercase', '大写字母', '科学计数法E和十六进制字母大写'],
    ['ios::showpos', '显示正号', '正数前面显示+号'],
    ['ios::scientific', '科学计数法', '以指数形式显示浮点数'],
    ['ios::fixed', '固定小数点', '以固定小数点形式显示浮点数'],
    ['ios::unitbuf', '每次输出后刷新', '每次<<操作后自动刷新缓冲区'],
    ['ios::stdio', '同步C标准IO', '与C的stdout/stderr同步（默认启用）'],
]
flags_table = create_table(doc, flags_headers, flags_rows)
elements_to_insert.append(flags_table)

# Code example
elements_to_insert.append(make_normal_paragraph(doc, ''))
flags_ex1 = make_normal_paragraph(doc, '// 使用格式标志设置输出格式')
elements_to_insert.append(flags_ex1)
flags_ex2 = make_normal_paragraph(doc, 'int inum = 255;')
elements_to_insert.append(flags_ex2)
flags_ex3 = make_normal_paragraph(doc, 'cout.flags(ios::oct | ios::showbase);   // 设置八进制+显示前缀')
elements_to_insert.append(flags_ex3)
flags_ex4 = make_normal_paragraph(doc, 'cout << inum;                           // 输出: 0377')
elements_to_insert.append(flags_ex4)
flags_ex5 = make_normal_paragraph(doc, '// 取消某标志: cout.unsetf(ios::skipws);  取消跳过空白')
elements_to_insert.append(flags_ex5)

# === 4.5 输入流成员函数 ===
elements_to_insert.append(make_normal_paragraph(doc, ''))
p45_header = make_bold_paragraph(doc, '4.5 输入流成员函数')
elements_to_insert.append(p45_header)
elements_to_insert.append(make_normal_paragraph(doc, ''))

input_headers = ['函数', '功能', '示例与说明']
input_rows = [
    ['cin.get()', '读取一个字符（含空格和换行）', 'char c = cin.get();  // 不跳过空白'],
    ['cin.get(ch)', '读取一个字符存入 ch', 'char ch; cin.get(ch);'],
    ['cin.getline(buf, n)', '读取一行到字符数组（最多 n-1 字符）', 'char buf[100]; cin.getline(buf, 100);'],
    ['cin.getline(buf, n, delim)', '读取到指定分隔符为止', 'cin.getline(buf, 100, \',\');  // 遇到逗号停止'],
    ['cin.ignore(n, delim)', '忽略输入中最多 n 个字符（直到 delim）', 'cin.ignore(100, \'\\n\');  // 忽略剩余一行'],
    ['cin.gcount()', '返回上次未格式化输入读取的字符数', 'int count = cin.gcount();'],
    ['cin.eof()', '判断是否到达文件末尾', 'if (cin.eof()) { ... }'],
    ['cin.peek()', '查看下一个字符但不提取', 'char next = cin.peek();'],
    ['cin.putback(ch)', '将字符 ch 放回输入流', 'cin.putback(\'A\');'],
    ['cin.read(buf, n)', '读取 n 个字符到缓冲区', 'char buf[100]; cin.read(buf, 50);'],
]
input_table = create_table(doc, input_headers, input_rows)
elements_to_insert.append(input_table)

# Code example
elements_to_insert.append(make_normal_paragraph(doc, ''))
input_ex1 = make_normal_paragraph(doc, '// 读取整行字符串（含空格）')
elements_to_insert.append(input_ex1)
input_ex2 = make_normal_paragraph(doc, 'string line;')
elements_to_insert.append(input_ex2)
input_ex3 = make_normal_paragraph(doc, 'getline(cin, line);  // 读取一行到 string（推荐方式）')
elements_to_insert.append(input_ex3)
input_ex4 = make_normal_paragraph(doc, '// 注意：先用 cin >> 再用 getline 时，需先清除换行符')
elements_to_insert.append(input_ex4)
input_ex5 = make_normal_paragraph(doc, 'cin >> n; cin.ignore();  // 忽略 >> 残留的换行符')
elements_to_insert.append(input_ex5)

# === 4.6 文件打开模式 ===
elements_to_insert.append(make_normal_paragraph(doc, ''))
p46_header = make_bold_paragraph(doc, '4.6 文件打开模式详解')
elements_to_insert.append(p46_header)
elements_to_insert.append(make_normal_paragraph(doc, ''))

mode_headers = ['模式', '值', '说明']
mode_rows = [
    ['ios::in', '0x01', '打开文件用于读取（ifstream 默认）'],
    ['ios::out', '0x02', '打开文件用于写入（ofstream 默认）'],
    ['ios::ate', '0x04', '打开时定位到文件末尾（可随后移动指针）'],
    ['ios::app', '0x08', '追加模式：每次写入都自动定位到文件末尾'],
    ['ios::trunc', '0x10', '打开时清空已有文件内容（ofstream 默认）'],
    ['ios::binary', '0x80', '以二进制方式打开（否则为文本模式）'],
    ['ios::nocreate', '0x20', '若文件不存在则打开失败（不创建新文件）'],
    ['ios::noreplace', '0x40', '若文件已存在则打开失败（不覆盖已有文件）'],
]
mode_table = create_table(doc, mode_headers, mode_rows)
elements_to_insert.append(mode_table)

# Code example
elements_to_insert.append(make_normal_paragraph(doc, ''))
mode_ex1 = make_normal_paragraph(doc, '// 组合使用文件打开模式')
elements_to_insert.append(mode_ex1)
mode_ex2 = make_normal_paragraph(doc, 'fstream file;')
elements_to_insert.append(mode_ex2)
mode_ex3 = make_normal_paragraph(doc, 'file.open("data.txt", ios::in | ios::out | ios::binary);  // 二进制读写')
elements_to_insert.append(mode_ex3)
mode_ex4 = make_normal_paragraph(doc, 'file.open("log.txt", ios::out | ios::app);  // 追加写入，不覆盖原有内容')
elements_to_insert.append(mode_ex4)
mode_ex5 = make_normal_paragraph(doc, '// 检查文件是否成功打开')
elements_to_insert.append(mode_ex5)
mode_ex6 = make_normal_paragraph(doc, 'ifstream fin("input.txt", ios::in | ios::nocreate);')
elements_to_insert.append(mode_ex6)
mode_ex7 = make_normal_paragraph(doc, 'if (!fin) { cout << "文件不存在!" << endl; exit(1); }')
elements_to_insert.append(mode_ex7)

# === 4.7 流状态检查 ===
elements_to_insert.append(make_normal_paragraph(doc, ''))
p47_header = make_bold_paragraph(doc, '4.7 流状态检查')
elements_to_insert.append(p47_header)
elements_to_insert.append(make_normal_paragraph(doc, ''))

state_headers = ['函数', '功能', '说明']
state_rows = [
    ['good()', '流状态正常', '最近操作成功，流可用'],
    ['fail()', '操作失败（可恢复）', '如：类型不匹配，格式错误'],
    ['eof()', '到达文件末尾', '读取操作遇到文件结束符'],
    ['bad()', '严重错误（不可恢复）', '如：磁盘读写错误，流被破坏'],
    ['clear()', '清除错误状态', '恢复流为可用状态：cin.clear()'],
    ['clear(0)', '清除所有标志位', '将流状态重置为正常'],
    ['rdstate()', '返回当前流状态', '返回状态位组合，用于调试'],
]
state_table = create_table(doc, state_headers, state_rows)
elements_to_insert.append(state_table)

# Code example
elements_to_insert.append(make_normal_paragraph(doc, ''))
state_ex1 = make_normal_paragraph(doc, '// 检查流状态典型用法')
elements_to_insert.append(state_ex1)
state_ex2 = make_normal_paragraph(doc, 'ifstream fin("data.txt");')
elements_to_insert.append(state_ex2)
state_ex3 = make_normal_paragraph(doc, 'if (!fin.good()) { cout << "打开失败!" << endl; exit(1); }')
elements_to_insert.append(state_ex3)
state_ex4 = make_normal_paragraph(doc, 'while (!fin.eof()) {  // 循环读取直到文件末尾')
elements_to_insert.append(state_ex4)
state_ex5 = make_normal_paragraph(doc, '    fin >> x;')
elements_to_insert.append(state_ex5)
state_ex6 = make_normal_paragraph(doc, '    if (!fin.eof()) { /* 处理 x */ }  // 注意：eof() 在读完后才为 true')
elements_to_insert.append(state_ex6)
state_ex7 = make_normal_paragraph(doc, '}')
elements_to_insert.append(state_ex7)

# === 4.8 二进制文件读写 ===
elements_to_insert.append(make_normal_paragraph(doc, ''))
p48_header = make_bold_paragraph(doc, '4.8 二进制文件读写')
elements_to_insert.append(p48_header)
elements_to_insert.append(make_normal_paragraph(doc, ''))

bin_headers = ['函数', '功能', '示例']
bin_rows = [
    ['read(char* buf, int n)', '从输入流读取 n 字节', 'fin.read((char*)&data, sizeof(data));'],
    ['write(const char* buf, int n)', '向输出流写入 n 字节', 'fout.write((char*)&data, sizeof(data));'],
    ['seekg(offset, dir)', '设置输入流指针位置', 'fin.seekg(0, ios::beg);  // 移到文件开头'],
    ['seekp(offset, dir)', '设置输出流指针位置', 'fout.seekp(0, ios::end);  // 移到文件末尾'],
    ['tellg()', '返回输入流当前指针位置', 'streampos pos = fin.tellg();'],
    ['tellp()', '返回输出流当前指针位置', 'streampos pos = fout.tellp();'],
]
bin_table = create_table(doc, bin_headers, bin_rows)
elements_to_insert.append(bin_table)

# Positioning directions
elements_to_insert.append(make_normal_paragraph(doc, ''))
dir_ex1 = make_normal_paragraph(doc, '// 指针定位方向（用于 seekg/seekp 的第二个参数）：')
elements_to_insert.append(dir_ex1)
dir_ex2 = make_normal_paragraph(doc, '// ios::beg — 从文件开头偏移；ios::cur — 从当前位置偏移；ios::end — 从文件末尾偏移')
elements_to_insert.append(dir_ex2)

# === 4.9 自定义类型的 I/O 运算符重载 ===
elements_to_insert.append(make_normal_paragraph(doc, ''))
p49_header = make_bold_paragraph(doc, '4.9 自定义类型的 I/O 运算符重载')
elements_to_insert.append(p49_header)
elements_to_insert.append(make_normal_paragraph(doc, ''))

overload_text = make_normal_paragraph(doc, '可以为自定义类重载 >> 和 << 运算符，使其能像内置类型一样输入输出。')
elements_to_insert.append(overload_text)
elements_to_insert.append(make_normal_paragraph(doc, ''))

overload_code1 = make_normal_paragraph(doc, 'class Complex {')
elements_to_insert.append(overload_code1)
overload_code2 = make_normal_paragraph(doc, '    double Real, Image;')
elements_to_insert.append(overload_code2)
overload_code3 = make_normal_paragraph(doc, 'public:')
elements_to_insert.append(overload_code3)
overload_code4 = make_normal_paragraph(doc, '    Complex(double r = 0.0, double i = 0.0) : Real(r), Image(i) {}')
elements_to_insert.append(overload_code4)
overload_code5 = make_normal_paragraph(doc, '    // 声明友元函数，使其能访问私有成员')
elements_to_insert.append(overload_code5)
overload_code6 = make_normal_paragraph(doc, '    friend istream& operator>>(istream& is, Complex& c);')
elements_to_insert.append(overload_code6)
overload_code7 = make_normal_paragraph(doc, '    friend ostream& operator<<(ostream& os, const Complex& c);')
elements_to_insert.append(overload_code7)
overload_code8 = make_normal_paragraph(doc, '};')
elements_to_insert.append(overload_code8)
overload_code9 = make_normal_paragraph(doc, '')
elements_to_insert.append(overload_code9)
overload_code10 = make_normal_paragraph(doc, '// 重载 >> 运算符：从输入流读取 Complex 对象')
elements_to_insert.append(overload_code10)
overload_code11 = make_normal_paragraph(doc, 'istream& operator>>(istream& is, Complex& c) {')
elements_to_insert.append(overload_code11)
overload_code12 = make_normal_paragraph(doc, '    is >> c.Real >> c.Image;')
elements_to_insert.append(overload_code12)
overload_code13 = make_normal_paragraph(doc, '    return is;  // 必须返回流引用以支持链式调用')
elements_to_insert.append(overload_code13)
overload_code14 = make_normal_paragraph(doc, '}')
elements_to_insert.append(overload_code14)
overload_code15 = make_normal_paragraph(doc, '')
elements_to_insert.append(overload_code15)
overload_code16 = make_normal_paragraph(doc, '// 重载 << 运算符：向输出流写入 Complex 对象')
elements_to_insert.append(overload_code16)
overload_code17 = make_normal_paragraph(doc, 'ostream& operator<<(ostream& os, const Complex& c) {')
elements_to_insert.append(overload_code17)
overload_code18 = make_normal_paragraph(doc, '    os << c.Real << "+" << c.Image << "i";')
elements_to_insert.append(overload_code18)
overload_code19 = make_normal_paragraph(doc, '    return os;')
elements_to_insert.append(overload_code19)
overload_code20 = make_normal_paragraph(doc, '}')
elements_to_insert.append(overload_code20)
overload_code21 = make_normal_paragraph(doc, '')
elements_to_insert.append(overload_code21)
overload_code22 = make_normal_paragraph(doc, '// 使用示例:')
elements_to_insert.append(overload_code22)
overload_code23 = make_normal_paragraph(doc, 'Complex c1(3, 4), c2, c3;')
elements_to_insert.append(overload_code23)
overload_code24 = make_normal_paragraph(doc, 'cin >> c2 >> c3;          // 连续输入两个复数')
elements_to_insert.append(overload_code24)
overload_code25 = make_normal_paragraph(doc, 'cout << c1 << endl;       // 输出: 3+4i')
elements_to_insert.append(overload_code25)

# === 4.10 文件复制示例 ===
elements_to_insert.append(make_normal_paragraph(doc, ''))
p410_header = make_bold_paragraph(doc, '4.10 文件复制完整示例')
elements_to_insert.append(p410_header)
elements_to_insert.append(make_normal_paragraph(doc, ''))

copy_code1 = make_normal_paragraph(doc, 'int main() {')
elements_to_insert.append(copy_code1)
copy_code2 = make_normal_paragraph(doc, '    ifstream sfile("input.txt");         // 打开源文件')
elements_to_insert.append(copy_code2)
copy_code3 = make_normal_paragraph(doc, '    ofstream dfile("output.txt");        // 打开目标文件')
elements_to_insert.append(copy_code3)
copy_code4 = make_normal_paragraph(doc, '    if (!sfile || !dfile) {              // 检查是否成功打开')
elements_to_insert.append(copy_code4)
copy_code5 = make_normal_paragraph(doc, '        cout << "文件打开失败!" << endl;')
elements_to_insert.append(copy_code5)
copy_code6 = make_normal_paragraph(doc, '        return 1;')
elements_to_insert.append(copy_code6)
copy_code7 = make_normal_paragraph(doc, '    }')
elements_to_insert.append(copy_code7)
copy_code8 = make_normal_paragraph(doc, '    char ch;')
elements_to_insert.append(copy_code8)
copy_code9 = make_normal_paragraph(doc, '    sfile.unsetf(ios::skipws);           // 不跳过空白字符')
elements_to_insert.append(copy_code9)
copy_code10 = make_normal_paragraph(doc, '    while (sfile >> ch) dfile << ch;     // 逐字符复制')
elements_to_insert.append(copy_code10)
copy_code11 = make_normal_paragraph(doc, '    sfile.close(); dfile.close();        // 关闭文件')
elements_to_insert.append(copy_code11)
copy_code12 = make_normal_paragraph(doc, '    return 0;')
elements_to_insert.append(copy_code12)
copy_code13 = make_normal_paragraph(doc, '}')
elements_to_insert.append(copy_code13)

# === 4.11 字符串流 ===
elements_to_insert.append(make_normal_paragraph(doc, ''))
p411_header = make_bold_paragraph(doc, '4.11 字符串流（需 #include <sstream>）')
elements_to_insert.append(p411_header)
elements_to_insert.append(make_normal_paragraph(doc, ''))

ss_headers = ['类', '功能', '示例']
ss_rows = [
    ['istringstream', '从字符串读取数据', 'istringstream iss("123 4.56 hello");'],
    ['ostringstream', '将数据写入字符串', 'ostringstream oss; oss << "result=" << 42;'],
    ['stringstream', '同时支持读写', 'stringstream ss; ss << 100; int x; ss >> x;'],
]
ss_table = create_table(doc, ss_headers, ss_rows)
elements_to_insert.append(ss_table)

elements_to_insert.append(make_normal_paragraph(doc, ''))
ss_ex1 = make_normal_paragraph(doc, '// 字符串流使用示例：字符串与数值转换')
elements_to_insert.append(ss_ex1)
ss_ex2 = make_normal_paragraph(doc, 'string s = "3.14 2.718";')
elements_to_insert.append(ss_ex2)
ss_ex3 = make_normal_paragraph(doc, 'istringstream iss(s);')
elements_to_insert.append(ss_ex3)
ss_ex4 = make_normal_paragraph(doc, 'double a, b;')
elements_to_insert.append(ss_ex4)
ss_ex5 = make_normal_paragraph(doc, 'iss >> a >> b;  // a = 3.14, b = 2.718（从字符串中解析数值）')
elements_to_insert.append(ss_ex5)
ss_ex6 = make_normal_paragraph(doc, '')
elements_to_insert.append(ss_ex6)
ss_ex7 = make_normal_paragraph(doc, '// 格式化输出到字符串')
elements_to_insert.append(ss_ex7)
ss_ex8 = make_normal_paragraph(doc, 'ostringstream oss;')
elements_to_insert.append(ss_ex8)
ss_ex9 = make_normal_paragraph(doc, 'oss << fixed << setprecision(2) << 3.14159;')
elements_to_insert.append(ss_ex9)
ss_ex10 = make_normal_paragraph(doc, 'string result = oss.str();  // result = "3.14"')
elements_to_insert.append(ss_ex10)

# === Now insert all elements before P43 (Section 5) ===
print(f"Inserting {len(elements_to_insert)} elements before Section 5...")

# Get the parent (body) and index of P43
body = doc.element.body
children = list(body)
p43_index = children.index(p43_element)

# Insert elements in reverse order so they end up in the right position
for i, elem in enumerate(reversed(elements_to_insert)):
    body.insert(p43_index, elem._element)

# Save the modified document
output_path = r'C:\Users\lenovo\Desktop\常用运算符和数学函数_含示例解释.docx'
doc.save(output_path)
print(f"Document saved successfully to: {output_path}")
print(f"Added {len(elements_to_insert)} elements")
