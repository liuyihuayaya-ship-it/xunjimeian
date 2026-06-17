#!/usr/bin/env python
# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn

docx_path = r'C:\Users\lenovo\Desktop\常用运算符和数学函数_含示例解释.docx'
doc = Document(docx_path)

# Check formatting of key paragraphs
print("=== Original Section Headers (baseline) ===")
for idx in [2, 18, 38]:
    p = doc.paragraphs[idx]
    r = p.runs[0] if p.runs else None
    if r:
        rPr = r._r.find(qn('w:rPr'))
        rFonts = rPr.find(qn('w:rFonts')) if rPr is not None else None
        color = rPr.find(qn('w:color')) if rPr is not None else None
        ascii_f = rFonts.get(qn('w:ascii')) if rFonts is not None else 'N/A'
        east_f = rFonts.get(qn('w:eastAsia')) if rFonts is not None else 'N/A'
        c_val = color.get(qn('w:val')) if color is not None else 'N/A'
        print(f"P{idx} '{p.text[:50]}': bold={r.bold}, size={r.font.size}, "
              f"ascii={ascii_f}, eastAsia={east_f}, color={c_val}")
        # spacing
        pPr_p = p._p.find(qn('w:pPr'))
        sp = pPr_p.find(qn('w:spacing')) if pPr_p is not None else None
        if sp is not None:
            print(f"  spacing: before={sp.get(qn('w:before'))}, after={sp.get(qn('w:after'))}")
        else:
            print(f"  spacing: None")

print("\n=== New Section Headers (4.3, 4.4, etc.) ===")
for idx, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text and text.startswith('4.') and not text.startswith('4.2'):
        r = p.runs[0] if p.runs else None
        if r:
            rPr = r._r.find(qn('w:rPr'))
            rFonts = rPr.find(qn('w:rFonts')) if rPr is not None else None
            color = rPr.find(qn('w:color')) if rPr is not None else None
            ascii_f = rFonts.get(qn('w:ascii')) if rFonts is not None else 'N/A'
            east_f = rFonts.get(qn('w:eastAsia')) if rFonts is not None else 'N/A'
            c_val = color.get(qn('w:val')) if color is not None else 'N/A'
            print(f"P{idx} '{text[:60]}': bold={r.bold}, size={r.font.size}, "
                  f"ascii={ascii_f}, eastAsia={east_f}, color={c_val}")
            pPr_p = p._p.find(qn('w:pPr'))
            sp = pPr_p.find(qn('w:spacing')) if pPr_p is not None else None
            if sp is not None:
                print(f"  spacing: before={sp.get(qn('w:before'))}, after={sp.get(qn('w:after'))}")
            else:
                print(f"  spacing: None")

print("\n=== New Code Paragraphs ===")
for idx in [45, 46, 47]:  # first few code lines after 4.3 section
    p = doc.paragraphs[idx]
    r = p.runs[0] if p.runs else None
    if r:
        rPr = r._r.find(qn('w:rPr'))
        rFonts = rPr.find(qn('w:rFonts')) if rPr is not None else None
        color = rPr.find(qn('w:color')) if rPr is not None else None
        ascii_f = rFonts.get(qn('w:ascii')) if rFonts is not None else 'N/A'
        east_f = rFonts.get(qn('w:eastAsia')) if rFonts is not None else 'N/A'
        c_val = color.get(qn('w:val')) if color is not None else 'N/A'
        print(f"P{idx} '{p.text[:70]}': bold={r.bold}, size={r.font.size}, "
              f"ascii={ascii_f}, eastAsia={east_f}, color={c_val}")
        pPr_p = p._p.find(qn('w:pPr'))
        sp = pPr_p.find(qn('w:spacing')) if pPr_p is not None else None
        if sp is not None:
            print(f"  spacing: before={sp.get(qn('w:before'))}, after={sp.get(qn('w:after'))}")
        else:
            print(f"  spacing: None")

print("\n=== New Tables (header and data cells) ===")
# Table 9 should be 4.3 I/O manipulators (first new table)
for ti in [9, 10, 11]:
    if ti < len(doc.tables):
        t = doc.tables[ti]
        print(f"\nTable {ti} ({len(t.rows)} rows):")
        # Header
        hdr = t.rows[0].cells[0]
        for r in hdr.paragraphs[0].runs:
            rPr = r._r.find(qn('w:rPr'))
            rFonts = rPr.find(qn('w:rFonts')) if rPr is not None else None
            ascii_f = rFonts.get(qn('w:ascii')) if rFonts is not None else 'N/A'
            east_f = rFonts.get(qn('w:eastAsia')) if rFonts is not None else 'N/A'
            print(f"  Header: text='{r.text[:20]}', bold={r.bold}, size={r.font.size}, "
                  f"ascii={ascii_f}, eastAsia={east_f}")
        # Data
        data = t.rows[1].cells[0]
        for r in data.paragraphs[0].runs:
            rPr = r._r.find(qn('w:rPr'))
            rFonts = rPr.find(qn('w:rFonts')) if rPr is not None else None
            ascii_f = rFonts.get(qn('w:ascii')) if rFonts is not None else 'N/A'
            east_f = rFonts.get(qn('w:eastAsia')) if rFonts is not None else 'N/A'
            print(f"  Data:   text='{r.text[:20]}', bold={r.bold}, size={r.font.size}, "
                  f"ascii={ascii_f}, eastAsia={east_f}")

# Compare with original tables
print("\n=== Original Tables (for comparison) ===")
for ti in [0, 1]:
    t = doc.tables[ti]
    print(f"\nOriginal Table {ti}:")
    hdr = t.rows[0].cells[0]
    for r in hdr.paragraphs[0].runs:
        rPr = r._r.find(qn('w:rPr'))
        rFonts = rPr.find(qn('w:rFonts')) if rPr is not None else None
        ascii_f = rFonts.get(qn('w:ascii')) if rFonts is not None else 'N/A'
        east_f = rFonts.get(qn('w:eastAsia')) if rFonts is not None else 'N/A'
        print(f"  Header: text='{r.text[:20]}', bold={r.bold}, size={r.font.size}, "
              f"ascii={ascii_f}, eastAsia={east_f}")
    data = t.rows[1].cells[0]
    for r in data.paragraphs[0].runs:
        rPr = r._r.find(qn('w:rPr'))
        rFonts = rPr.find(qn('w:rFonts')) if rPr is not None else None
        ascii_f = rFonts.get(qn('w:ascii')) if rFonts is not None else 'N/A'
        east_f = rFonts.get(qn('w:eastAsia')) if rFonts is not None else 'N/A'
        print(f"  Data:   text='{r.text[:20]}', bold={r.bold}, size={r.font.size}, "
              f"ascii={ascii_f}, eastAsia={east_f}")
