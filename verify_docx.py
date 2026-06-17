#!/usr/bin/env python
# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document

docx_path = r'C:\Users\lenovo\Desktop\常用运算符和数学函数_含示例解释.docx'
doc = Document(docx_path)

print("=== Updated document paragraphs ===")
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text:
        bold = any(r.bold for r in p.runs if r.bold)
        marker = "**" if bold else "  "
        print(f"P{i} {marker}: {text[:120]}")

print(f"\n=== Tables: {len(doc.tables)} ===")
for i, t in enumerate(doc.tables):
    header = [c.text[:30] for c in t.rows[0].cells]
    print(f"Table {i} ({len(t.rows)} rows): {header}")
