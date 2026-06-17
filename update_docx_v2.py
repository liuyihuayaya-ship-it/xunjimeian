#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Update docx with proper formatting matching original document."""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

docx_path = r'C:\Users\lenovo\Desktop\常用运算符和数学函数_含示例解释.docx'
doc = Document(docx_path)

FONT_NAME = '微软雅黑'
SECTION_COLOR = '003366'   # dark blue for major sections (四、五...)
SUBSEC_COLOR = '336633'    # dark green for subsections (4.1, 4.2...)
SECTION_SIZE = Pt(11)      # 139700 EMU
CODE_SIZE = Pt(10)         # 127000 EMU
TABLE_SIZE = Pt(9)         # 114300 EMU

def set_run_font(run, font_name=FONT_NAME, font_size=CODE_SIZE, bold=False, color=None):
    """Configure a run to match original document formatting."""
    run.font.name = font_name
    run.font.size = font_size
    run.bold = bold
    rPr = run._r.get_or_add_rPr()
    # Remove existing rFonts
    for old_rf in rPr.findall(qn('w:rFonts')):
        rPr.remove(old_rf)
    # Add proper rFonts with all three font slots
    rFonts = parse_xml(
        f'<w:rFonts {nsdecls("w")} '
        f'w:ascii="{font_name}" w:hAnsi="{font_name}" w:eastAsia="{font_name}"/>'
    )
    rPr.insert(0, rFonts)
    # Remove existing color
    for old_color in rPr.findall(qn('w:color')):
        rPr.remove(old_color)
    # Add color if specified
    if color:
        color_el = parse_xml(f'<w:color {nsdecls("w")} w:val="{color}"/>')
        rPr.append(color_el)

def remove_paragraph_spacing(p):
    """Remove space_before and space_after from paragraph (matching original)."""
    pPr = p._p.get_or_add_pPr()
    for old_sp in pPr.findall(qn('w:spacing')):
        pPr.remove(old_sp)

def make_section_header(doc, text):
    """Section header: bold, 11pt, 微软雅黑, dark blue, no spacing."""
    p = doc.add_paragraph()
    remove_paragraph_spacing(p)
    run = p.add_run(text)
    set_run_font(run, font_size=SECTION_SIZE, bold=True, color=SECTION_COLOR)
    return p

def make_subsection_header(doc, text):
    """Subsection header: bold, 11pt, 微软雅黑, dark green, no spacing."""
    p = doc.add_paragraph()
    remove_paragraph_spacing(p)
    run = p.add_run(text)
    set_run_font(run, font_size=SECTION_SIZE, bold=True, color=SUBSEC_COLOR)
    return p

def make_code_para(doc, text):
    """Code line: not bold, 10pt, 微软雅黑, no spacing."""
    p = doc.add_paragraph()
    remove_paragraph_spacing(p)
    if text:  # only add run if non-empty
        run = p.add_run(text)
        set_run_font(run, font_size=CODE_SIZE, bold=False, color=None)
    return p

def make_empty_line(doc):
    """Empty line like original spacing between sections."""
    return make_code_para(doc, '')

def create_table(doc, headers, rows):
    """Create table matching original: Table Grid, header bold 9pt, data 9pt."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        remove_paragraph_spacing(p)
        run = p.add_run(h)
        set_run_font(run, font_size=TABLE_SIZE, bold=True, color=None)

    # Data rows
    for r, row_data in enumerate(rows):
        for c, cell_text in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            remove_paragraph_spacing(p)
            run = p.add_run(str(cell_text))
            set_run_font(run, font_size=TABLE_SIZE, bold=False, color=None)

    return table


# === Collect all new elements ===
elements = []

# ===== 4.3 I/O 操纵符 =====
elements.append(make_subsection_header(doc, '4.3 I/O 操纵符（需 #include <iomanip>）'))
elements.append(make_empty_line(doc))

manip_rows = [
    ['setw(n)', '设置输出宽度为 n 个字符', 'cout << setw(10) << x;'],
    ['setfill(c)', '设置填充字符为 c', 'cout << setfill(\'*\') << setw(8) << x;'],
    ['setprecision(n)', '设置浮点数精度为 n 位', 'cout << setprecision(6) << pi;'],
    ['left', '输出左对齐', 'cout << left << setw(10) << x;'],
    ['right', '输出右对齐（默认）', 'cout << right << setw(10) << x;'],
    ['fixed', '固定小数点方式输出浮点数', 'cout << fixed << 3.14159;  // 3.141590'],
    ['scientific', '科学计数法方式输出浮点数', 'cout << scientific << 3.14;  // 3.140000e+000'],
    ['hex', '以十六进制输出整数', 'cout << hex << 255;  // ff'],
    ['oct', '以八进制输出整数', 'cout << oct << 255;  // 377'],
    ['dec', '以十进制输出整数（默认）', 'cout << dec << 255;  // 255'],
    ['showbase', '显示进制前缀（0x 或 0）', 'cout << showbase << hex << 255;  // 0xff'],
    ['showpoint', '始终显示小数点', 'cout << showpoint << 5.0;  // 5.00000'],
    ['endl', '换行并刷新输出缓冲区', 'cout << "Hello" << endl;'],
    ['\\n', '仅换行，不刷新（比 endl 快）', 'cout << "Hello\\n";'],
]
elements.append(create_table(doc, ['操纵符', '功能', '示例'], manip_rows))

elements.append(make_empty_line(doc))
elements.append(make_code_para(doc, '// 格式控制示例：'))
elements.append(make_code_para(doc, 'int inum = 255;'))
elements.append(make_code_para(doc, 'cout << hex << showbase << inum;           // 输出: 0xff（十六进制带前缀）'))
elements.append(make_code_para(doc, 'cout << right << setw(10) << setfill(\'*\') << inum;  // 输出: *******255'))
elements.append(make_code_para(doc, ''))
elements.append(make_code_para(doc, 'double pi = 3.1415926535;'))
elements.append(make_code_para(doc, 'cout << fixed << setprecision(6) << pi;   // 输出: 3.141593（固定6位小数）'))
elements.append(make_code_para(doc, 'cout << scientific << setprecision(9) << pi;  // 输出: 3.141592654e+000'))

# ===== 4.4 格式控制标志 =====
elements.append(make_empty_line(doc))
elements.append(make_subsection_header(doc, '4.4 格式控制标志（ios::flags）'))
elements.append(make_empty_line(doc))

flags_rows = [
    ['ios::skipws', '跳过空白字符', '输入时跳过前导空格、制表符（默认启用）'],
    ['ios::left', '左对齐输出', '在设定宽度内靠左显示'],
    ['ios::right', '右对齐输出', '在设定宽度内靠右显示（默认）'],
    ['ios::internal', '内部对齐', '符号靠左，数值靠右，中间填充字符'],
    ['ios::dec', '十进制', '整数以十进制显示（默认）'],
    ['ios::oct', '八进制', '整数以八进制显示'],
    ['ios::hex', '十六进制', '整数以十六进制显示'],
    ['ios::showbase', '显示进制前缀', '八进制加前缀 0，十六进制加前缀 0x'],
    ['ios::showpoint', '显示小数点', '浮点数始终显示小数点'],
    ['ios::uppercase', '大写字母', '科学计数法 E 和十六进制字母 A-F 使用大写'],
    ['ios::showpos', '显示正号', '正数前面显示 + 号'],
    ['ios::scientific', '科学计数法', '以指数形式显示浮点数'],
    ['ios::fixed', '固定小数点', '以固定小数点形式显示浮点数'],
    ['ios::unitbuf', '每次输出后刷新', '每次 << 操作后自动刷新缓冲区'],
    ['ios::stdio', '同步 C 标准 I/O', '与 stdout / stderr 同步刷新（默认启用）'],
]
elements.append(create_table(doc, ['标志', '功能', '说明'], flags_rows))

elements.append(make_empty_line(doc))
elements.append(make_code_para(doc, '// 使用格式标志设置输出格式：'))
elements.append(make_code_para(doc, 'int inum = 255;'))
elements.append(make_code_para(doc, 'cout.flags(ios::oct | ios::showbase);  // 设置为八进制 + 显示前缀'))
elements.append(make_code_para(doc, 'cout << inum;                          // 输出: 0377'))
elements.append(make_code_para(doc, ''))
elements.append(make_code_para(doc, '// 取消某个标志：'))
elements.append(make_code_para(doc, 'cout.unsetf(ios::skipws);  // 取消跳过空白，之后可读取空格'))

# ===== 4.5 输入流成员函数 =====
elements.append(make_empty_line(doc))
elements.append(make_subsection_header(doc, '4.5 输入流成员函数'))
elements.append(make_empty_line(doc))

input_rows = [
    ['cin.get()', '读取一个字符（含空格和换行）', 'char c = cin.get();  // 不跳过任何空白字符'],
    ['cin.get(ch)', '读取一个字符存入 ch', 'char ch; cin.get(ch);'],
    ['cin.getline(buf, n)', '读取一行到字符数组', 'char buf[100]; cin.getline(buf, 100);'],
    ['cin.getline(buf, n, delim)', '读取到指定分隔符为止', 'cin.getline(buf, 100, \',\');  // 遇到逗号停止'],
    ['cin.ignore(n, delim)', '忽略输入中最多 n 个字符', 'cin.ignore(100, \'\\n\');  // 忽略剩余一行'],
    ['cin.ignore()', '忽略一个字符（默认换行）', 'cin.ignore();  // 常用于清除残留的换行符'],
    ['cin.gcount()', '返回上次未格式化输入读取的字符数', 'int n = cin.gcount();'],
    ['cin.eof()', '判断是否到达文件末尾', 'if (cin.eof()) { ... }'],
    ['cin.peek()', '查看下一个字符但不提取', 'char next = cin.peek();'],
    ['cin.putback(ch)', '将字符 ch 放回输入流', 'cin.putback(\'A\');'],
    ['cin.read(buf, n)', '读取 n 个字符到缓冲区', 'char buf[100]; cin.read(buf, 50);'],
]
elements.append(create_table(doc, ['函数', '功能', '示例与说明'], input_rows))

elements.append(make_empty_line(doc))
elements.append(make_code_para(doc, '// 读取整行字符串（含空格）——推荐方式：'))
elements.append(make_code_para(doc, 'string line;'))
elements.append(make_code_para(doc, 'getline(cin, line);  // 读取一整行到 string'))
elements.append(make_code_para(doc, ''))
elements.append(make_code_para(doc, '// 注意：先用 cin >> 再用 getline 时，需先清除残留换行符：'))
elements.append(make_code_para(doc, 'cin >> n;'))
elements.append(make_code_para(doc, 'cin.ignore();  // 忽略 >> 之后残留的换行符'))
elements.append(make_code_para(doc, 'getline(cin, line);  // 现在才能正确读取下一行'))

# ===== 4.6 文件打开模式详解 =====
elements.append(make_empty_line(doc))
elements.append(make_subsection_header(doc, '4.6 文件打开模式详解'))
elements.append(make_empty_line(doc))

mode_rows = [
    ['ios::in', '0x01', '打开文件用于读取（ifstream 默认包含）'],
    ['ios::out', '0x02', '打开文件用于写入（ofstream 默认包含）'],
    ['ios::ate', '0x04', '打开时定位到文件末尾，之后可移动读写指针'],
    ['ios::app', '0x08', '追加模式：每次写入前自动定位到文件末尾'],
    ['ios::trunc', '0x10', '打开时清空已有文件内容（ofstream 默认包含）'],
    ['ios::binary', '0x80', '以二进制方式打开（默认以文本方式打开）'],
    ['ios::nocreate', '0x20', '若文件不存在则打开失败，不创建新文件'],
    ['ios::noreplace', '0x40', '若文件已存在则打开失败，不覆盖已有文件'],
]
elements.append(create_table(doc, ['模式', '值', '说明'], mode_rows))

elements.append(make_empty_line(doc))
elements.append(make_code_para(doc, '// 组合使用打开模式：'))
elements.append(make_code_para(doc, 'fstream file;'))
elements.append(make_code_para(doc, 'file.open("data.txt", ios::in | ios::out | ios::binary);  // 二进制读写'))
elements.append(make_code_para(doc, 'file.open("log.txt", ios::out | ios::app);     // 追加写入，不覆盖'))
elements.append(make_code_para(doc, ''))
elements.append(make_code_para(doc, '// 以不创建新文件的方式打开：'))
elements.append(make_code_para(doc, 'ifstream fin("input.txt", ios::in | ios::nocreate);'))
elements.append(make_code_para(doc, 'if (!fin) {  // 文件不存在时打开失败'))
elements.append(make_code_para(doc, '    cout << "文件不存在！" << endl;'))
elements.append(make_code_para(doc, '    exit(1);'))
elements.append(make_code_para(doc, '}'))

# ===== 4.7 流状态检查 =====
elements.append(make_empty_line(doc))
elements.append(make_subsection_header(doc, '4.7 流状态检查'))
elements.append(make_empty_line(doc))

state_rows = [
    ['good()', '流状态正常，无任何错误', '最近一次 I/O 操作成功完成'],
    ['fail()', '操作失败（可恢复）', '如：类型不匹配、格式错误（读入 int 却输入字母）'],
    ['eof()', '到达文件末尾', '读取操作遇到文件结束符 EOF'],
    ['bad()', '发生严重错误（不可恢复）', '如：磁盘读写错误，流缓冲区被破坏'],
    ['clear()', '清除所有错误标志', '恢复流为可用状态：cin.clear();'],
    ['clear(0)', '清除所有标志位（重置）', '将流状态完全重置为正常状态'],
    ['rdstate()', '返回当前流的状态标志', '返回各状态位的组合值，可用于调试'],
]
elements.append(create_table(doc, ['函数', '功能', '说明'], state_rows))

elements.append(make_empty_line(doc))
elements.append(make_code_para(doc, '// 流状态检查典型用法：'))
elements.append(make_code_para(doc, 'ifstream fin("data.txt");'))
elements.append(make_code_para(doc, 'if (!fin.good()) { cout << "文件打开失败！" << endl; exit(1); }'))
elements.append(make_code_para(doc, ''))
elements.append(make_code_para(doc, '// 循环读取直到文件末尾：'))
elements.append(make_code_para(doc, 'while (!fin.eof()) {'))
elements.append(make_code_para(doc, '    fin >> x;'))
elements.append(make_code_para(doc, '    if (!fin.eof()) {  // 注意：eof() 在读取失败后才变为 true'))
elements.append(make_code_para(doc, '        /* 成功读取，处理 x */'))
elements.append(make_code_para(doc, '    }'))
elements.append(make_code_para(doc, '}'))

# ===== 4.8 二进制文件读写 =====
elements.append(make_empty_line(doc))
elements.append(make_subsection_header(doc, '4.8 二进制文件读写'))
elements.append(make_empty_line(doc))

bin_rows = [
    ['read(char* buf, int n)', '从输入流读取 n 个字节', 'fin.read((char*)&data, sizeof(data));'],
    ['write(const char* buf, int n)', '向输出流写入 n 个字节', 'fout.write((char*)&data, sizeof(data));'],
    ['seekg(offset, dir)', '移动输入流指针', 'fin.seekg(0, ios::beg);  // 移到文件开头'],
    ['seekp(offset, dir)', '移动输出流指针', 'fout.seekp(0, ios::end);  // 移到文件末尾'],
    ['tellg()', '获取输入流当前指针位置', 'streampos pos = fin.tellg();'],
    ['tellp()', '获取输出流当前指针位置', 'streampos pos = fout.tellp();'],
]
elements.append(create_table(doc, ['函数', '功能', '示例'], bin_rows))

elements.append(make_empty_line(doc))
elements.append(make_code_para(doc, '// 指针定位方向（seekg / seekp 的第二个参数）：'))
elements.append(make_code_para(doc, '// ios::beg — 从文件开头偏移；ios::cur — 从当前位置偏移；ios::end — 从文件末尾偏移'))

# ===== 4.9 自定义类型的 I/O 运算符重载 =====
elements.append(make_empty_line(doc))
elements.append(make_subsection_header(doc, '4.9 自定义类型的 I/O 运算符重载'))
elements.append(make_empty_line(doc))
elements.append(make_code_para(doc, '可以为自定义类重载 >> 和 << 运算符，使其能像内置类型一样输入输出。'))
elements.append(make_empty_line(doc))
elements.append(make_code_para(doc, 'class Complex {'))
elements.append(make_code_para(doc, '    double Real, Image;'))
elements.append(make_code_para(doc, 'public:'))
elements.append(make_code_para(doc, '    Complex(double r = 0.0, double i = 0.0) : Real(r), Image(i) {}'))
elements.append(make_code_para(doc, ''))
elements.append(make_code_para(doc, '    // 声明友元函数，使其能访问私有成员'))
elements.append(make_code_para(doc, '    friend istream& operator>>(istream& is, Complex& c);'))
elements.append(make_code_para(doc, '    friend ostream& operator<<(ostream& os, const Complex& c);'))
elements.append(make_code_para(doc, '};'))
elements.append(make_empty_line(doc))
elements.append(make_code_para(doc, '// 重载 >> 运算符：从输入流读取 Complex 对象'))
elements.append(make_code_para(doc, 'istream& operator>>(istream& is, Complex& c) {'))
elements.append(make_code_para(doc, '    is >> c.Real >> c.Image;'))
elements.append(make_code_para(doc, '    return is;  // 必须返回流引用，以支持链式调用'))
elements.append(make_code_para(doc, '}'))
elements.append(make_empty_line(doc))
elements.append(make_code_para(doc, '// 重载 << 运算符：向输出流写入 Complex 对象'))
elements.append(make_code_para(doc, 'ostream& operator<<(ostream& os, const Complex& c) {'))
elements.append(make_code_para(doc, '    os << c.Real << "+" << c.Image << "i";'))
elements.append(make_code_para(doc, '    return os;'))
elements.append(make_code_para(doc, '}'))
elements.append(make_empty_line(doc))
elements.append(make_code_para(doc, '// 使用示例：'))
elements.append(make_code_para(doc, 'Complex c1(3, 4), c2, c3;'))
elements.append(make_code_para(doc, 'cin >> c2 >> c3;          // 连续输入两个复数'))
elements.append(make_code_para(doc, 'cout << c1 << endl;       // 输出: 3+4i'))

# ===== 4.10 文件复制完整示例 =====
elements.append(make_empty_line(doc))
elements.append(make_subsection_header(doc, '4.10 文件复制完整示例'))
elements.append(make_empty_line(doc))
elements.append(make_code_para(doc, 'int main() {'))
elements.append(make_code_para(doc, '    ifstream sfile("input.txt");            // 打开源文件'))
elements.append(make_code_para(doc, '    ofstream dfile("output.txt");           // 打开目标文件'))
elements.append(make_code_para(doc, ''))
elements.append(make_code_para(doc, '    if (!sfile || !dfile) {                 // 检查是否成功打开'))
elements.append(make_code_para(doc, '        cout << "文件打开失败！" << endl;'))
elements.append(make_code_para(doc, '        return 1;'))
elements.append(make_code_para(doc, '    }'))
elements.append(make_code_para(doc, ''))
elements.append(make_code_para(doc, '    char ch;'))
elements.append(make_code_para(doc, '    sfile.unsetf(ios::skipws);              // 不跳过空白字符'))
elements.append(make_code_para(doc, '    while (sfile >> ch) dfile << ch;        // 逐字符复制'))
elements.append(make_code_para(doc, ''))
elements.append(make_code_para(doc, '    sfile.close(); dfile.close();           // 关闭文件'))
elements.append(make_code_para(doc, '    return 0;'))
elements.append(make_code_para(doc, '}'))

# ===== 4.11 字符串流 =====
elements.append(make_empty_line(doc))
elements.append(make_subsection_header(doc, '4.11 字符串流（需 #include <sstream>）'))
elements.append(make_empty_line(doc))

ss_rows = [
    ['istringstream', '从字符串中读取数据', 'istringstream iss("123 4.56 hello");'],
    ['ostringstream', '将数据格式化写入字符串', 'ostringstream oss; oss << "result = " << 42;'],
    ['stringstream', '同时支持从字符串读和向字符串写', 'stringstream ss;  // 双向操作'],
]
elements.append(create_table(doc, ['类', '功能', '示例'], ss_rows))

elements.append(make_empty_line(doc))
elements.append(make_code_para(doc, '// 字符串流示例：字符串与数值的转换'))
elements.append(make_code_para(doc, 'string s = "3.14 2.718";'))
elements.append(make_code_para(doc, 'istringstream iss(s);'))
elements.append(make_code_para(doc, 'double a, b;'))
elements.append(make_code_para(doc, 'iss >> a >> b;  // a = 3.14, b = 2.718（从字符串解析数值）'))
elements.append(make_empty_line(doc))
elements.append(make_code_para(doc, '// 格式化输出到字符串：'))
elements.append(make_code_para(doc, 'ostringstream oss;'))
elements.append(make_code_para(doc, 'oss << fixed << setprecision(2) << 3.14159;'))
elements.append(make_code_para(doc, 'string result = oss.str();  // result = "3.14"'))


# === Now insert all before Section 5 ===
print(f"Total elements to insert: {len(elements)}")

# Find Section 5 paragraph
body = doc.element.body
p5_element = None
for p in body.findall(qn('w:p')):
    texts = p.findall(qn('w:t'))
    combined = ''.join(t.text or '' for t in texts)
    if '五、常用数学函数' in combined:
        p5_element = p
        break

if p5_element is None:
    print("ERROR: Could not find Section 5!")
    sys.exit(1)

children = list(body)
p5_index = children.index(p5_element)
print(f"Section 5 at index: {p5_index}")

# Find and remove old added elements (between 4.2 area and Section 5)
# We need to identify elements that were added in the first run and remove them
# The content between table for 4.2 and P5 should be removed first

# Find the table that follows 4.2 (Table 8 originally, now whatever follows)
# Strategy: find P41 (4.2 header), then find the next table, then remove everything between that table and P5

# Find P41 (4.2 header)
p42_header = None
for p in body.findall(qn('w:p')):
    texts = p.findall(qn('w:t'))
    combined = ''.join(t.text or '' for t in texts)
    if '4.2' in combined and '文件读写' in combined:
        p42_header = p
        break

if p42_header:
    print(f"Found 4.2 header")

# Remove old inserted content: find the table immediately after 4.2 paragraph
# then everything after that table up to section 5 is old inserted content

# Simple approach: find all elements between the 4.2 section (P41 + its table) and Section 5
# The structure should be: ... P41(4.2) ... Table(文件读写) ... [OLD_INSERTED_CONTENT] ... P5(五)
# We need to remove [OLD_INSERTED_CONTENT]

# Walk backwards from P5 to find and remove old content
# We'll look for the last table that belongs to the original document (the 4.2 table)
# then everything after it until P5 is old content

# Find all elements between p42_header and p5_element
p42_idx = children.index(p42_header)
print(f"4.2 header at index: {p42_idx}")

# Find the 4.2 table (it should be right after 4.2 header)
# It will be the next table element after p42_header
table_4_2 = None
for elem in children[p42_idx+1:p5_index]:
    if elem.tag == qn('w:tbl'):
        table_4_2 = elem
        break

if table_4_2:
    table_4_2_idx = children.index(table_4_2)
    print(f"4.2 table at index: {table_4_2_idx}")
    # Remove everything between table_4_2+1 and p5_element
    to_remove = []
    for elem in children[table_4_2_idx+1:p5_index]:
        to_remove.append(elem)

    print(f"Removing {len(to_remove)} old inserted elements...")
    for elem in to_remove:
        body.remove(elem)
else:
    # Fallback: remove everything between p42_header+1 and p5_element
    to_remove = []
    for elem in children[p42_idx+1:p5_index]:
        to_remove.append(elem)
    print(f"Fallback: Removing {len(to_remove)} elements between 4.2 header and Section 5...")
    for elem in to_remove:
        body.remove(elem)

# Re-find p5_element after removal
children = list(body)
for p in body.findall(qn('w:p')):
    texts = p.findall(qn('w:t'))
    combined = ''.join(t.text or '' for t in texts)
    if '五、常用数学函数' in combined:
        p5_element = p
        break
p5_index = children.index(p5_element)
print(f"Section 5 now at index: {p5_index}")

# Insert new elements before P5
for elem in reversed(elements):
    body.insert(p5_index, elem._element)

# Save
doc.save(docx_path)
print(f"\nSaved successfully! Added {len(elements)} elements with matching format.")
