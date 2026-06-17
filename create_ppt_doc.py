#!/usr/bin/env python3
"""Create a C++ OOP knowledge points Word document from PPT folder PDFs."""

import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# ── Styles ──
STYLE_CODE = 'CodeBlock'
if STYLE_CODE not in [s.name for s in doc.styles]:
    s = doc.styles.add_style(STYLE_CODE, WD_STYLE_TYPE.PARAGRAPH)
    s.font.name = 'Consolas'; s.font.size = Pt(8.5)
    s.paragraph_format.space_before = Pt(1)
    s.paragraph_format.space_after = Pt(1)
    s.paragraph_format.left_indent = Cm(0.5)
    rPr = s.element.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr'); s.element.append(rPr)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '微软雅黑'); rPr.insert(0, rFonts)

NOTE_STYLE = 'NoteBlock'
if NOTE_STYLE not in [s.name for s in doc.styles]:
    s = doc.styles.add_style(NOTE_STYLE, WD_STYLE_TYPE.PARAGRAPH)
    s.font.name = '微软雅黑'; s.font.size = Pt(9)
    s.font.color.rgb = RGBColor(0xCC, 0x33, 0x00)
    s.paragraph_format.left_indent = Cm(0.5)

def _set_font(run, name='微软雅黑', size=Pt(10)):
    run.font.name = name; run.font.size = size
    rPr = run._r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr'); run._r.insert(0, rPr)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rPr.insert(0, rFonts)

def H(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs: _set_font(r, size={1:Pt(14),2:Pt(12),3:Pt(10.5),4:Pt(10)}.get(level,Pt(10)))

def T(text, bold=False, size=Pt(10)):
    p = doc.add_paragraph()
    r = p.add_run(text)
    _set_font(r, size=size); r.bold = bold
    return p

def C(code_text):
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph(style=STYLE_CODE)
        r = p.add_run(line)
        r.font.name = 'Consolas'; r.font.size = Pt(8.5)
    doc.add_paragraph().paragraph_format.space_before = Pt(1)

def N(text):
    p = doc.add_paragraph(style=NOTE_STYLE)
    p.add_run('⚠ ' + text)

def B(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    r = p.add_run(text); _set_font(r, size=Pt(9.5))

def TB(headers, rows):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            for r in p.runs: r.font.size = Pt(9); r.font.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(8.5)
    doc.add_paragraph()

# ============================================================
# TITLE
# ============================================================
title = doc.add_heading('C++ 面向对象程序设计 知识点总汇', level=0)
for r in title.runs: _set_font(r, size=Pt(18))
T('来源：桌面 PPT 文件夹（6 个 PDF 课件）', size=Pt(9))
T('内容：类与对象 · 模板 · 动态内存分配 · 继承与多态 · 流类库与输入输出', size=Pt(9))
doc.add_paragraph()

# ============================================================
# 第4章 类与对象
# ============================================================
H('第4章  类与对象', 1)

H('4.1 类与对象的基本概念', 2)

H('类的定义', 3)
T('类是一种用户自定义数据类型。定义类时系统不分配内存，创建对象时才分配。')
C('''class 类名 {
private:      // 私有成员（类外不能访问，默认访问权限）
    成员表1;
public:       // 公有成员（类外可访问）
    成员表2;
protected:    // 保护成员（类外不能访问，派生类可访问）
    成员表3;
};''')

H('成员函数定义', 3)
C('''返回值类型 类名::函数名(参数表) { …… }   // "::" 是作用域解析运算符''')
T('内联函数：用 inline 修饰，免去调用开销。inline 只是建议，由编译器决定是否执行。')
T('对象访问：对象名.公有成员 或 指向对象的指针->公有成员。只有公有成员才能在对象外访问。')

H('4.2 构造函数', 2)

H('构造函数特征', 3)
B('函数名与类名相同，无返回类型')
B('对象建立时自动调用，生存期只调用一次')
B('可以重载（多个不同参数的构造函数）')
B('若未定义任何构造函数，编译器自动提供默认无参构造函数：类名(void) { }')
N('一旦自定义了任何构造函数，系统不再提供默认无参构造函数！')

C('''class CGoods {
    char Name[21]; int Amount; float Price;
public:
    CGoods();                                        // 默认（无参）构造函数
    CGoods(char name[], int amount, float price);    // 三参数构造函数
    CGoods(char name[], float price);                // 两参数构造函数（重载）
};''')

H('带默认参数的构造函数', 3)
C('''Complex(double r = 0.0, double i = 0.0);  // 利用默认参数，一个替代两个构造函数''')

H('初始化列表', 3)
C('''类名::构造函数名(参数表): 成员1(初值1), 成员2(初值2), … { …… }
// 例：
Complex(double r, double i): Real(r), Image(i) { }''')

H('4.3 拷贝构造函数（重点：深拷贝 vs 浅拷贝）', 2)

H('拷贝构造函数基本概念', 3)
T('拷贝构造函数：用已有对象初始化新对象。参数必须是类对象的引用（const 类名&），否则会死循环。')
T('使用场合：')
B('① 用一个对象初始化另一个对象：CGoods g2(g1); 或 CGoods g2 = g1;')
B('② 函数形参是类对象（值传递时）')
B('③ 函数返回类对象时')
N('C++11 起：编译器优先通过 RVO/NRVO（返回值优化）消除临时对象。')
N('一旦自定义拷贝构造函数，默认拷贝构造函数就不再起作用。')

H('浅拷贝（Shallow Copy）— 问题所在', 3)
T('默认拷贝构造函数是值拷贝（逐成员复制）。当类中有指针成员时，两个对象的指针指向同一块内存，互不独立。')
T('后果：一个对象析构释放空间后，另一个对象的指针变成野指针，双重释放导致崩溃。')
C('''// 默认拷贝构造函数：浅拷贝
student s2 = s1;   // s2.pName 和 s1.pName 指向同一块堆内存！
// 析构时同一空间被 delete 两次 → 程序崩溃''')

H('深拷贝（Deep Copy）— 解决方案', 3)
T('当类中有指针成员时，必须实现深拷贝，让两个对象拥有各自独立的内存空间。')
T('深拷贝四要素（当对象有指针成员时，以下四个必须全部自定义）：')
B('① 重载构造函数 — 用 new 申请独立空间')
B('② 必须自定义拷贝构造函数 — 深拷贝，申请新空间并复制内容')
B('③ 重载析构函数 — 用 delete 释放空间')
B('④ 重载赋值运算符 = — 深拷贝 + 防止自赋值 + 先释放旧空间再分配新空间')

C('''// 深拷贝赋值运算符重载（完整版）
student& student::operator=(const student& st) {
    if (this == &st) return *this;           // ① 防止自赋值！
    if (pName != NULL) delete[] pName;        // ② 先释放旧空间
    if (st.pName) {                           // ③ 分配新空间并复制
        pName = new char[strlen(st.pName) + 1];
        strcpy_s(pName, strlen(st.pName) + 1, st.pName);
    } else pName = nullptr;
    return *this;
}''')
N('封装的好处：使用 string 代替 char*，不需要手动处理深拷贝和浅拷贝问题！')

H('4.4 析构函数', 2)
T('特征：函数名前加 ~，无返回类型，无参数。一个类只有一个析构函数，可以默认。')
T('对象注销时系统自动调用。顺序与构造函数相反：先析构自身，再析构成员对象。')
C('''~CGoods() { }   // 析构函数示例''')

H('封闭类（含对象成员的类）构造与析构顺序', 3)
C('''类名::构造函数名(参数总表): 对象成员1(参数1), 对象成员2(参数2), … { …… }''')
T('构造顺序：基类 → 成员对象（按声明顺序）→ 自身构造函数体')
T('析构顺序：自身 → 成员对象 → 基类（与构造完全相反）')
T('先建立的对象后撤销（栈结构）。')

H('完整商品类案例', 3)
C('''class CGoods {
    char Name[21]; int Amount; float Price; float Total_value;
public:
    CGoods();                                    // 默认构造函数
    CGoods(char[], int, float);                  // 三参数构造函数
    CGoods(char[], float);                       // 两参数构造函数
    void RegisterGoods(char[], int, float);      // 注册商品
    void CountTotal(void);                       // 计算总价
    void GetName(char[]);                        // 获取名称
    int GetAmount(void);                         // 获取数量
    float GetPrice(void);                        // 获取单价
    float GetTotal_value(void);                  // 获取总价
};''')

H('4.5 运算符重载', 2)

H('成员函数实现格式', 3)
C('''返回值类型 类名::operator运算符(参数表) { …… }''')
B('重载双目运算符：左操作数为调用对象(this)，右操作数为函数参数')
B('重载单目运算符：无参数（操作数是 this）')
B('前置 ++：类名::operator++()')
B('后置 ++：类名::operator++(int) — int 仅用于区分，不实际使用')

C('''// 复数类运算符重载示例
Complex operator+(const Complex&);     // 加法
Complex operator*(const Complex&);     // 乘法
Complex operator=(Complex);            // 赋值
Complex operator+=(Complex);           // 复合赋值

// 前置++：先加后用
Complex& operator++() { Real++; Image++; return *this; }

// 后置++：先用后加（int 参数仅用于区分）
Complex operator++(int) {
    Complex temp(*this);
    Real++; Image++;
    return temp;
}''')

H('赋值运算符重载优化', 3)
C('''// 返回引用（不调用拷贝构造函数），效率高
Complex& Complex::operator=(const Complex& c) {
    Real = c.Real; Image = c.Image;
    return *this;
}''')

H('this 指针', 3)
T('系统自动产生的隐藏指针，指向调用成员函数的对象本身。常用于返回自身引用：return *this;')
N('赋值运算符 = 重载必须为成员函数，不可为友元函数。')

H('禁止重载的运算符', 3)
T('? :  .  .*  ::  sizeof — 这五个运算符不能重载')

H('4.6 友元（friend）', 2)

H('友元函数', 3)
T('不是类的成员函数，是普通函数，但可以在类外访问类的私有成员。破坏了封装性，慎用。')
C('''class CPoint {
    double m_x, m_y;
    friend double Cal_Distance(const CPoint&, const CPoint&);
};''')
T('友元函数实现运算符重载：左操作数和右操作数都作为参数传入。')
C('''// 友元函数重载 +
friend Complex operator+(const Complex& c1, const Complex& c2);
// 使用：d + c 被解释为 operator+(d, c)

// 友元函数重载前置++
friend Complex operator++(Complex&);
// 友元函数重载后置++（参数必须是引用）
friend Complex operator++(Complex&, int);''')

H('友元类', 3)
C('''class CPoint { friend CLine; };  // CLine 可直接访问 CPoint 所有私有成员''')

H('4.7 静态成员（static）', 2)

H('静态数据成员', 3)
B('用 static 修饰，所有对象共享一个存储空间')
B('存放在全局变量区，具有静态生命期')
B('使用格式：类名::静态数据成员名')
B('static const 可在类内初始化；static 非常量必须在类外初始化')
N('不能在构造函数中初始化静态成员！')

H('静态成员函数', 3)
B('用 static 修饰，属于整个类，不属于某个对象')
B('使用格式：类名::函数名(参数)')
B('静态成员函数没有 this 指针，不能访问非静态成员')

H('4.8 结构（struct）与名空间（namespace）', 2)
T('C++ 中 struct 成员默认为 public（class 默认为 private）。')
T('同结构类型变量间可整体赋值；结构可嵌套定义。')
T('namespace 用于避免命名冲突。C++ 标准库在名空间 std 中。using namespace std; // 引入全部')

H('4.9 标准 C++ string 类', 2)
C('''string str1;                // 空串
string str2("OK");          // C 字符串初始化
string str3(str2);          // 拷贝构造
str1 = str2;                // 赋值
str1 += str2;               // 连接
str1 = str2 + str3;         // 拼接
str.at(i);                  // 边界检查访问
str[i];                     // 无边界检查访问（快但不安全）
str.c_str();                // 转为 C 风格字符串
str.substr(pos, len);       // 提取子串
str.empty();                // 是否为空
str.length();               // 字符串长度
str.insert(pos, str2);      // 在 pos 位置插入 str2''')

# ============================================================
# 第6章 模板
# ============================================================
H('第6章  模板', 1)

H('6.1 模板概念', 2)
T('模板 = 泛型编程 / 参数化程序设计。将数据类型参数化，实现与类型无关的通用代码，达到代码重用。')

H('6.2 函数模板', 2)

H('定义格式', 3)
C('''template <模板参数表>
返回类型 函数名(形式参数表) { …… }''')

H('函数模板 vs 模板函数', 3)
B('函数模板：模板定义本身（蓝图）')
B('模板函数：函数模板经实例化生成的具体函数（实例）')
B('模板实参推演：隐式（自动推导类型）和显式（max<double>(d, e) 手动指定）')
N('模板函数不允许自动类型转换，普通函数允许。')

H('示例', 3)
C('''template <typename T>
T max(T a, T b) { return a > b ? a : b; }

int m = max(3, 7);             // T 推导为 int
double d = max(3.14, 2.72);    // T 推导为 double
double e = max<double>(3, 7.5);// 显式指定 T = double''')

H('函数模板匹配原则', 3)
B('非模板函数和同名函数模板可同时存在')
B('优先调用非模板函数（精确匹配优先）')
B('可通过空模板实参列表强制用模板：max<>(a, b)')
N('函数模板声明和实现通常不能分离到 .h 和 .cpp（除非显式实例化）。')

H('矩阵运算模板示例', 3)
C('''template <typename T1, typename T2>
void inverse(T1* mat1, T2* mat2, int a, int b);

template <typename T1, typename T2>
void multi(T1* mat1, T2* mat2, T2* result, int a, int b, int c);''')

H('6.3 类模板', 2)

H('定义格式', 3)
C('''// 类模板定义
template <模板参数表> class 类名 { …… };

// 成员函数在类外定义
template <模板参数表>
返回类型 类名<模板参数名表>::成员函数名(形参表) { …… }''')

H('非类型模板参数', 3)
C('''template <typename T, int MaxSize>
class CVector {
    T VectValue[MaxSize];
    int RealSize;
    // ...
};
// 使用：CVector<int, 100> vec;''')
N('float、double 不能做非类型模板参数！只能是整数类型。')

H('关键性质', 3)
B('类模板的成员函数均为函数模板')
B('类模板实例化时成员函数不自动实例化，只有被调用时才实例化')
B('类模板的声明和实现通常放在一个 .h 文件里')

H('顺序表类模板 CSeqList 提供的方法', 3)
T('Find（查找）、IsIn（判断存在）、Insert（插入）、Remove（删除）、Next（后继）、Prior（前驱）、IsEmpty/IsFull（判空/判满）、BinarySearch（二分查找）、BubbleSort（冒泡排序）、InsertSort（插入排序）、BinaryInsertSort（对半插入排序）、重载下标运算符[]。')

H('6.4 排序算法复杂度比较', 2)
TB(
    ['算法', '平均时间', '最好情况', '最差情况', '稳定性'],
    [
        ['选择排序', 'O(n²)', 'n(n-1)/2', 'n(n-1)/2', '不稳定'],
        ['冒泡排序', 'O(n²)', 'n-1', 'n(n-1)/2', '稳定'],
        ['插入排序', 'O(n²)', 'n-1', 'n(n-1)/2', '稳定'],
    ]
)

H('6.5 索引查找与指针数组', 2)
T('为避免大数据移动，通过索引（目录）实现查找和排序。利用指针数组操作数据，实现对指针排序而非数据移动。')

# ============================================================
# 第7章 动态内存分配
# ============================================================
H('第7章  动态内存分配', 1)

H('7.1 内存分区', 2)
TB(
    ['内存区域', '存放内容', '特点'],
    [
        ['栈区（Stack）', '局部变量、函数形参', '分配时不处理内存（随机值）；函数退出自动释放'],
        ['堆区（Heap）', '动态分配的数据（new 出来的）', '生命期由程序员控制，delete 后才释放'],
        ['全局/静态数据区', '全局变量、static 变量', '分配时初始化为全 0；程序结束时释放'],
        ['代码区', '函数体的二进制代码', '只读'],
    ]
)

H('7.2 new 和 delete', 2)
C('''// ——— 单个变量 ———
int* pi = new int(10);        // 申请并初始化为 10
int* pi1 = new int;           // 未初始化（值为随机值）
int* pi2 = new int();         // 初始化为 0
delete pi;                    // 释放单个变量

// ——— 动态数组 ———
double* pdd = new double[n];  // 申请 n 个 double 的数组
delete[] pdd;                 // 释放数组（[] 必不可少！）

// ——— 多维数组 ———
float (*bp)[20] = new float[30][20];  // 30 行 20 列
delete[] bp;''')

H('7.3 动态内存管理注意事项', 2)
B('① 分配失败检查：new 失败返回 nullptr，应检查 if (!p) { exit(1); }')
B('② 内存泄漏：int* p = new int(3); p = new int(5); // 第一个 new 的空间没 delete 就丢了！')
B('③ 重复释放非常危险：delete p; 后再 delete p 会导致未定义行为')
B('④ 动态变量的生命期不依赖于建立它的作用域，只有 delete 后才结束')
B('⑤ 动态数组中的元素并未被初始化')

H('7.4 深拷贝与浅拷贝（动态内存场景下的重点）', 2)

H('问题回顾', 3)
T('当对象有指针成员指向动态内存时，默认拷贝构造函数是值拷贝（浅拷贝），导致两个对象的指针指向同一块堆内存。')
T('后果：析构时同一空间被 delete 两次 → 程序崩溃。')

H('解决方案：深拷贝四要素', 3)
B('① 重载构造函数：用 new 申请独立堆空间')
B('② 必须自定义拷贝构造函数：为新对象申请独立空间并复制内容')
B('③ 重载析构函数：用 delete/delete[] 释放动态空间')
B('④ 重载赋值运算符：先释放旧空间，再分配新空间并复制，防止自赋值')

C('''// 深拷贝赋值运算符（完整版）
student& student::operator=(const student& st) {
    if (this == &st) return *this;           // ① 防止自赋值
    if (pName != NULL) delete[] pName;        // ② 先释放旧空间
    if (st.pName) {                           // ③ 分配新空间并复制
        pName = new char[strlen(st.pName) + 1];
        strcpy_s(pName, strlen(st.pName) + 1, st.pName);
    } else pName = nullptr;
    return *this;                              // ④ 返回自身引用
}''')
N('最佳实践：使用 string 代替 char*，不需要手动处理深拷贝和浅拷贝问题！')

H('动态顺序表模板（含深拷贝）', 3)
C('''template <typename T>
class Seqlist2 {
    T* slist;          // 指向动态数组的指针
    int maxsize;       // 最大容量
    int last;          // 最后一个元素的下标
public:
    Seqlist2();                           // 构造 — new 空间
    Seqlist2(const Seqlist2&);            // 深拷贝构造！
    ~Seqlist2() { delete[] slist; }       // 析构 — delete 空间
    // ... 其他操作
};''')

H('7.5 链表及其基本操作', 2)

H('链表概念', 3)
T('每个结点包含：数据域(info) + 指针域(link)。头结点可选，但可使操作统一。')

C('''struct node {
    int info;       // 数据域
    node* link;     // 指针域（指向下一个结点）
};''')

H('基本操作', 3)
C('''// 遍历链表
void print(node* head) {
    node* p = head;
    while (p != NULL) {
        cout << p->info << " ";
        p = p->link;
    }
}

// 查找
node* search(node* head, int key) {
    node* p = head;
    while (p != NULL && p->info != key)
        p = p->link;
    return p;   // 找到返回指针，没找到返回 NULL
}

// 插入（在结点 p 之后插入值为 e 的新结点）
void insert_after(node* p, int e) {
    node* newnode = new node;
    newnode->info = e;
    newnode->link = p->link;    // 新结点先指向后继
    p->link = newnode;           // 前驱再指向新结点（顺序不能反！）
}

// 删除 p 的后继结点
void del_after(node* p) {
    node* q = p->link;           // 保存待删结点
    p->link = q->link;           // 绕过待删结点
    delete q;                    // 释放空间
}''')

H('结点模板与链表模板', 3)
T('Node 类模板：info, link, InsertAfter(), RemoveAfter()')
T('List 类模板：head, tail, Find(), Length(), InsertFront/Rear/Order(), CreateNode(), DeleteNode(), MakeEmpty()')

# ============================================================
# 第8章 继承与多态
# ============================================================
H('第8章  继承与多态', 1)

H('8.1 继承的基本概念', 2)
B('基类（base class / 父类 / 超类）和派生类（derived class / 子类）')
B('派生类吸收基类除构造/析构外的全部成员（共性）+ 发展新成员（个性）')
B('同名覆盖（override）：派生类声明与基类成员同名的新成员，屏蔽基类同名成员')

H('派生类定义格式', 3)
C('''class 派生类名: 访问限定符 基类名1, …, 访问限定符 基类名n {
private:     成员表1;
public:      成员表2;
protected:   成员表3;
};''')

H('8.2 继承方式（访问控制）', 2)
TB(
    ['派生方式', '基类 public 成员', '基类 protected 成员', '基类 private 成员', '派生类外访问基类成员'],
    [
        ['public 继承', '→ public', '→ protected', '不可访问', 'public 成员可直接访问'],
        ['private 继承', '→ private', '→ private', '不可访问', '不可直接访问'],
        ['protected 继承', '→ protected', '→ protected', '不可访问', '不可直接访问'],
    ]
)

H('8.3 派生类的构造函数与析构函数', 2)

H('构造函数格式', 3)
C('''派生类名::派生类名(参数总表):
    基类名1(参数表1), …, 基类名n(参数表n),
    成员对象名1(参数表1), …, 成员对象名m(参数表m)
{
    // 新增成员初始化
}''')

H('调用顺序', 3)
T('构造：基类 → 成员对象（按声明次序）→ 派生类自身')
T('析构：派生类自身 → 成员对象 → 基类（与构造完全相反）')
N('若基类定义了带形参的构造函数，派生类必须定义构造函数！')
N('若派生类有新增指针成员，需要定义析构函数处理空间释放。')

H('8.4 赋值兼容规则（公有派生）', 2)
B('① 派生类对象可赋值给基类对象：Person p; Student s; p = s;（反之不行！）')
B('② 派生类对象地址可赋给基类指针：Person* p = &s;（但只能访问继承来的成员）')
B('③ 派生类对象可初始化基类引用：Person& ref = s;')

H('8.5 多重继承与虚继承', 2)

H('多重继承同名冲突', 3)
T('若多个基类有同名成员，用作用域运算符解决：obj.A::print(); obj.B::print();')

H('虚继承（virtual）— 解决菱形继承', 3)
T('菱形继承问题：B 和 C 继承 A，D 同时继承 B 和 C，导致 A 的成员在 D 中有两份。')
T('解决方案：B 和 C 虚继承 A，D 中只有一份 A 的成员。')
C('''class B: virtual public A { ... };
class C: virtual public A { ... };
class D: public B, public C {
    D(……): B(…), C(…), A(…) { … }  // 必须单独调用虚基类 A 的构造函数！
};''')

H('8.6 多态性与虚函数', 2)

H('多态分类', 3)
B('编译时多态（静态联编）：函数重载、运算符重载。编译时即确定调用哪个函数。')
B('运行时多态（动态联编）：虚函数。运行时根据实际对象类型确定调用哪个函数。')

H('虚函数定义', 3)
C('''virtual 返回类型 函数名(参数表) { … };
// 声明时加 virtual，类外定义时不加 virtual
// 派生类中覆盖(override)虚函数时，函数头必须完全相同''')

H('实现运行时多态两个必要条件', 3)
B('① 定义虚函数')
B('② 使用基类类型的指针或引用指向派生类对象，并通过该指针/引用调用虚函数')

C('''// === 经典多态示例 ===
class Animal {
public:
    virtual void speak() { cout << "动物叫声" << endl; }
    virtual ~Animal() {}  // ★ 虚析构函数（非常重要！）
};

class Dog: public Animal {
public:
    void speak() override { cout << "汪汪" << endl; }
};

class Cat: public Animal {
public:
    void speak() override { cout << "喵喵" << endl; }
};

int main() {
    Animal* p;           // 基类指针
    Dog dog; Cat cat;

    p = &dog;
    p->speak();          // 输出 "汪汪" ← 调用 Dog::speak()

    p = &cat;
    p->speak();          // 输出 "喵喵" ← 调用 Cat::speak()

    // 同一指针 p，不同对象调用不同函数 → 多态！
}''')

H('虚函数关键要点', 3)
B('派生类中覆盖虚函数时，函数头（返回类型、函数名、参数表）必须完全相同，否则是重载而非覆盖')
B('只有成员函数可为虚函数')
B('构造函数不能为虚函数')
B('析构函数可以为虚函数，且当用基类指针删除派生类对象时，必须将析构函数声明为虚函数！')
N('如果不用虚析构函数：delete 基类指针时只调用基类析构函数，派生类部分的内存泄漏！')
B('动态联编（基类指针/引用 → 虚函数）vs 静态联编（对象名.成员 的方式调用）')

H('纯虚函数与抽象类', 3)
C('''virtual 返回类型 函数名(参数表) = 0;   // 纯虚函数''')
B('含有纯虚函数的类是抽象类，不能用来定义对象')
B('抽象类用于定义接口框架，供派生类实现')
B('如果派生类没有实现所有纯虚函数，派生类也是抽象类')

# ============================================================
# 第9章 流类库与输入输出
# ============================================================
H('第9章  流类库与输入输出', 1)

H('9.1 C++ 基本流类体系', 2)
T('C++ 在标准库中包含 I/O 流类库，没有输入/输出语句。')
T('提取操作 >>：从流中取得数据。插入操作 <<：向流中添加数据。')
T('四个全局流对象：cin（标准输入/键盘）、cout（标准输出/显示器）、cerr、clog。')
T('头文件：<iostream>（标准I/O）、<fstream>（文件I/O）。')

H('流类层次', 3)
C('''basic_ios  ← 管理流缓冲区、I/O格式、错误处理
  ├── basic_istream  → cin
  ├── basic_ostream  → cout, cerr, clog
  ├── basic_ifstream    （文件输入流）
  ├── basic_ofstream    （文件输出流）
  ├── basic_iostream    （双向流）
  └── basic_fstream     （文件双向流）''')

H('9.2 I/O 格式控制', 2)

H('格式标志（常用）', 3)
TB(
    ['标志', '含义', '标志', '含义'],
    [
        ['skipws', '跳过空白', 'left', '左对齐'],
        ['right', '右对齐', 'dec', '十进制'],
        ['oct', '八进制', 'hex', '十六进制'],
        ['showbase', '显示数制前缀', 'showpoint', '必带小数点'],
        ['uppercase', '大写十六进制', 'showpos', '正数显示 +'],
        ['scientific', '科学计数法', 'fixed', '定点数形式'],
    ]
)

C('''// — 成员函数方式 —
cout.width(10);                      // 设置域宽
cout.fill('*');                      // 设置填充字符
cout.precision(6);                   // 设置精度
cout.flags(ios::oct | ios::showbase);// 八进制带基数

// — 操纵符方式（需 #include <iomanip>）—
cout << left << setw(8) << setfill('*');     // 左对齐，域宽8，填充*
cout << scientific << setprecision(9);       // 科学计数法，9位精度
cout << hex << showbase << uppercase << 255; // 输出 0XFF
cout << fixed << setprecision(2) << 3.14159; // 输出 3.14''')

H('9.3 标准设备 I/O', 2)
T('输入相关：cin.get()（含空白）、cin.getline()、cin.ignore()、cin.gcount()、cin.eof()')
T('状态字：当 I/O 操作出错时，必须 clear() 流对象才能继续使用。')

H('9.4 重载提取 >> 和插入 << 运算符', 2)
N('必须定义为友元函数！返回值为流引用，保证可连续使用 >> 或 <<。')
C('''// >> 重载格式
friend istream& operator>>(istream& is, 类名& obj);
// << 重载格式
friend ostream& operator<<(ostream& os, const 类名& obj);

// === 复数类完整示例 ===
class Complex {
    double Real, Image;
public:
    friend istream& operator>>(istream& is, Complex& c) {
        is >> c.Real >> c.Image;
        return is;
    }
    friend ostream& operator<<(ostream& os, const Complex& c) {
        os << c.Real << "+" << c.Image << "i";
        return os;   // 返回流引用，支持连续使用
    }
};''')

H('9.5 文件 I/O', 2)

H('使用文件三步骤', 3)
T('① 打开文件：建立流对象与磁盘文件的关联')
T('② 读写文件：提取 >> 或插入 << 数据')
T('③ 关闭文件：断开关联（析构时也会自动关闭）')

H('文件打开方式', 3)
TB(
    ['打开方式', '含义'],
    [
        ['ios::in', '打开用于读取'],
        ['ios::out', '打开用于写入（覆盖）'],
        ['ios::app', '在文件末尾追加写入'],
        ['ios::trunc', '打开时清空文件内容'],
        ['ios::binary', '以二进制方式打开'],
        ['ios::ate', '打开时文件指针定位到文件末尾'],
        ['ios::nocreate', '文件不存在则打开失败'],
        ['ios::noreplace', '文件已存在则打开失败'],
    ]
)

H('文本文件读写', 3)
C('''// ——— 逐字符复制文件 ———
ifstream sfile("test_in.txt");
ofstream dfile("test_out.txt");
sfile.unsetf(ios::skipws);    // 不跳过空白字符
char ch;
while (sfile >> ch)
    dfile << ch;
sfile.close(); dfile.close();

// ——— 按行复制 ———
char buf[100];
while (sfile.getline(buf, 100)) {
    if (sfile.gcount() < 100)
        dfile << buf << '\\n';
    else
        dfile << buf;
}''')

H('二进制文件读写', 3)
C('''// 读取指定字节数
istream& istream::read(char* buf, int size);
// 写入指定字节数
ostream& ostream::write(const char* buf, int size);

// 示例
fin.read((char*)&obj, sizeof(obj));
fout.write((char*)&obj, sizeof(obj));''')

H('9.6 文件与对象 — 面向对象程序固定框架', 2)
T('构造函数：打开文件 → 读取数据 → 创建对象')
T('析构函数：保存数据到文件 → 关闭文件 → 撤销对象')
C('''template <typename T>
class Array {
    T* elements;
    int Subscript, maxSize;
    fstream datafile;     // 文件流作为成员
public:
    Array(int = 10);      // 构造函数中打开文件并读入数据
    ~Array();             // 析构函数中保存数据并关闭文件
    void ordinsert(T&);   // 有序插入
};''')

# ============================================================
# 附录
# ============================================================
H('附录：关键概念索引（Ctrl+F 搜索直达）', 1)

TB(
    ['搜索关键词', '所在章节', '说明'],
    [
        ['深拷贝', '4.3', '拷贝构造函数，指针成员独立内存复制'],
        ['浅拷贝', '4.3', '默认逐成员复制，指针共享同一内存'],
        ['拷贝构造函数', '4.3', '用已有对象初始化新对象的构造函数'],
        ['构造函数', '4.2', '对象创建时自动调用的初始化函数'],
        ['析构函数', '4.4', '对象销毁时自动调用的清理函数'],
        ['运算符重载', '4.5', '自定义运算符对类对象的操作行为'],
        ['友元', '4.6', 'friend，可访问类私有成员的非成员函数/类'],
        ['静态成员', '4.7', 'static，所有对象共享的成员'],
        ['this 指针', '4.5', '指向调用对象自身的隐藏指针'],
        ['初始化列表', '4.2', '构造函数中初始化成员的高效方式'],
        ['string 类', '4.10', 'C++ 标准字符串类，自动管理内存'],
        ['函数模板', '6.2', '类型参数化的函数蓝图'],
        ['类模板', '6.3', '类型参数化的类蓝图'],
        ['排序算法', '6.4', '选择/冒泡/插入排序及复杂度分析'],
        ['new/delete', '7.2', '动态内存的申请与释放运算符'],
        ['内存泄漏', '7.3', '动态分配后未释放导致的内存浪费'],
        ['内存分区', '7.1', '栈区/堆区/全局数据区/代码区'],
        ['链表', '7.5', '结点的插入、删除、遍历等操作'],
        ['深拷贝赋值', '7.4', '重载 = 运算符，防止自赋值和内存泄漏'],
        ['继承', '8.1', '派生类从基类获得成员'],
        ['继承方式', '8.2', 'public/private/protected 继承'],
        ['赋值兼容', '8.4', '派生类对象可赋值给基类对象（公有继承）'],
        ['虚继承', '8.5', 'virtual 继承解决菱形继承问题'],
        ['多态', '8.6', '同一接口不同行为'],
        ['虚函数', '8.6', 'virtual 成员函数，支持运行时多态'],
        ['虚析构函数', '8.6', '基类析构函数声明为 virtual'],
        ['纯虚函数', '8.6', 'virtual ... = 0，定义接口'],
        ['抽象类', '8.6', '含纯虚函数的类，不能实例化'],
        ['override', '8.6', '派生类中覆盖基类虚函数'],
        ['动态联编', '8.6', '运行时确定调用哪个虚函数'],
        ['流类库', '9.1', 'iostream/fstream 体系'],
        ['重载>>和<<', '9.4', '自定义类型的流输入输出'],
        ['文件I/O', '9.5', '文本文件和二进制文件的读写'],
        ['格式控制', '9.2', 'setw/setprecision/进制控制等'],
    ]
)

doc.add_paragraph()
T('— 全文完 —', bold=True, size=Pt(11))

# ── Save ──
out_path = r'C:\Users\lenovo\Desktop\ppt.docx'
doc.save(out_path)
print(f'Done: {out_path}')
