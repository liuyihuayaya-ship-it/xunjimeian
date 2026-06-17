#!/usr/bin/env python
# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Emu
from docx.oxml.ns import qn

docx_path = r'C:\Users\lenovo\Desktop\常用运算符和数学函数_含示例解释.docx'
doc = Document(docx_path)

# Analyze original paragraphs (before the new content)
# Focus on P2 (section header), P5 (subsection), P20 (code), P38 (section header), P53 (tip)
# and their paragraph formatting

target_ps = [2, 5, 16, 20, 38, 39, 43, 53]

for idx in target_ps:
    p = doc.paragraphs[idx]
    print(f"\n=== P{idx}: '{p.text[:60]}' ===")
    pf = p.paragraph_format
    print(f"  alignment: {pf.alignment}")
    print(f"  space_before: {pf.space_before}")
    print(f"  space_after: {pf.space_after}")
    print(f"  line_spacing: {pf.line_spacing}")
    print(f"  first_line_indent: {pf.first_line_indent}")
    print(f"  style: {p.style.name if p.style else 'None'}")

    # Check paragraph XML for spacing
    pPr = p._p.find(qn('w:pPr'))
    if pPr is not None:
        spacing = pPr.find(qn('w:spacing'))
        if spacing is not None:
            print(f"  XML spacing: before={spacing.get(qn('w:before'))}, after={spacing.get(qn('w:after'))}, line={spacing.get(qn('w:line'))}, lineRule={spacing.get(qn('w:lineRule'))}")

        # Check for numbering/bullet
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            print(f"  Has numbering: {ET.tostring(numPr, encoding='unicode')}")

    for j, r in enumerate(p.runs):
        print(f"  Run {j}: text='{r.text[:40]}', bold={r.bold}, size={r.font.size}, name={r.font.name}")
        rPr = r._r.find(qn('w:rPr'))
        if rPr is not None:
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is not None:
                print(f"    RunFonts: ascii={rFonts.get(qn('w:ascii'))}, hAnsi={rFonts.get(qn('w:hAnsi'))}, eastAsia={rFonts.get(qn('w:eastAsia'))}")
            color = rPr.find(qn('w:color'))
            if color is not None:
                print(f"    Color: {color.get(qn('w:val'))}")

# Also analyze table formatting
print("\n\n=== Table Formatting ===")
for table_idx in [0, 1, 7]:  # data types, arithmetic ops, console I/O
    t = doc.tables[table_idx]
    print(f"\nTable {table_idx} style: {t.style.name}")
    # Check first row (header)
    hdr_cell = t.rows[0].cells[0]
    for j, p in enumerate(hdr_cell.paragraphs):
        for k, r in enumerate(p.runs):
            print(f"  Header cell[0] run {k}: text='{r.text[:30]}', bold={r.bold}, size={r.font.size}, name={r.font.name}")

    # Check data row
    if len(t.rows) > 1:
        data_cell = t.rows[1].cells[0]
        for j, p in enumerate(data_cell.paragraphs):
            for k, r in enumerate(p.runs):
                print(f"  Data cell[0] run {k}: text='{r.text[:30]}', bold={r.bold}, size={r.font.size}, name={r.font.name}")

# Check an original section header like P43 (五、常用数学函数) - which follows after tables
print("\n\n=== Checking spacing around Section 5 (P159 in updated doc) ===")
p159 = doc.paragraphs[159]
print(f"P159: '{p159.text[:80]}'")
pf = p159.paragraph_format
print(f"  space_before: {pf.space_before}, space_after: {pf.space_after}")
