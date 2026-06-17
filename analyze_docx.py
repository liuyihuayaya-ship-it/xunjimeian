#!/usr/bin/env python
# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import copy

docx_path = r'C:\Users\lenovo\Desktop\常用运算符和数学函数_含示例解释.docx'
doc = Document(docx_path)

# Print paragraph details with XML info
print("=== Paragraphs ===")
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text:
        # Check for bold runs
        bold_runs = []
        for r in p.runs:
            bold_runs.append(f"bold={r.bold}, size={r.font.size}, text='{r.text[:30]}'")
        runs_info = "; ".join(bold_runs[:3])
        print(f"P{i}: '{text[:80]}' | runs: [{runs_info}]")

print("\n=== Tables summary ===")
for i, t in enumerate(doc.tables):
    first_row = [c.text[:30] for c in t.rows[0].cells]
    print(f"Table {i}: {len(t.rows)} rows, header: {first_row}")

# Find exact insertion point
print("\n=== Insertion point search ===")
for i, p in enumerate(doc.paragraphs):
    if "五、" in p.text or "五." in p.text:
        print(f"Section 5 found at P{i}: '{p.text[:80]}'")
    if "4.2" in p.text:
        print(f"Section 4.2 found at P{i}: '{p.text[:80]}'")
