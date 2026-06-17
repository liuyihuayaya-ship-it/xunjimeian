#!/usr/bin/env python
"""Update the 常用运算符和数学函数_含示例解释.docx with ch9 content."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy
import os

docx_path = r'C:\Users\lenovo\Desktop\常用运算符和数学函数_含示例解释.docx'

doc = Document(docx_path)

# Helper: find the table that contains "文件读写" to know where to insert new content
# We'll find the paragraph after section 4.2 (文件读写 table) and insert before section 5

# Let's first print all paragraphs to understand the structure
print("=== Current document structure ===")
for i, p in enumerate(doc.paragraphs):
    style = p.style.name if p.style else "None"
    text_preview = p.text[:100] if p.text else "(empty)"
    print(f"P{i} [{style}]: {text_preview}")

print("\n=== Tables ===")
for i, t in enumerate(doc.tables):
    print(f"Table {i}: {len(t.rows)} rows x {len(t.columns)} cols")
    for j, row in enumerate(t.rows):
        cells = [c.text[:40] for c in row.cells]
        print(f"  Row {j}: {cells}")
