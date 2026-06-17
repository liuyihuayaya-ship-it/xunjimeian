from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'SimSun'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

def add_title(text, size=18):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.name = 'SimSun'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

def add_h(text, size=13):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.name = 'SimSun'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

def add_b(text, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = 'SimSun'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

def add_code(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.name = 'Courier New'

def add_pic(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.name = 'SimSun'
    r.font.color.rgb = RGBColor(192, 0, 0)
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

# ===== Cover =====
add_title('高等数学数学实验报告', 18)

p = doc.add_paragraph()
r = p.add_run('实验人员：院（系） ____________  学号 _______________  姓名 ______________\n实验地点：计算机中心机房')
r.font.size = Pt(11)
r.font.name = 'SimSun'
r._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
doc.add_paragraph()

# ===================================================================
# Experiment 1
# ===================================================================
add_h('实验一  空间曲线与曲面的绘制（习题1）', 15)

add_h('一、实验题目', 13)
add_b('作出各种标准二次曲面的图形。')

add_h('二、实验目的和意义', 13)
add_b('在高等数学课程中，我们学习了椭球面、椭圆抛物面、双曲抛物面、单叶双曲面、双叶双曲面、圆锥面等多种二次曲面。课本上的图片是静态的，只能从一个角度看，不太容易建立空间想象力。通过Mathematica软件绘制三维图形，可以旋转、放大缩小，从任意角度观察曲面的形状，帮助我直观地理解各种二次曲面的几何特征。')

add_h('三、计算公式', 13)
add_b('六种标准二次曲面的参数方程如下（参考实验讲义第6页）：')

add_b('1. 椭球面：x=a sin(u)cos(v), y=b sin(u)sin(v), z=c cos(u)\n    取a=1, b=1, c=2, u in [0,Pi], v in [0,2Pi]\n\n'
      '2. 椭圆抛物面：x=a u cos(v), y=b u sin(v), z=u^2\n    取a=1, b=2, u in [0,2], v in [0,2Pi]\n\n'
      '3. 双曲抛物面（马鞍面）：z = x^2 - y^2,  x,y in [-2,2]\n\n'
      '4. 单叶双曲面：x=a sec(u)cos(v), y=b sec(u)sin(v), z=c tan(u)\n    a=b=c=1, u in [-Pi/4,Pi/4], v in [0,2Pi]\n\n'
      '5. 双叶双曲面：x=a sqrt(1+u^2) cos(v), y=b sqrt(1+u^2) sin(v), z=c u\n    a=b=c=1, u in [-1.5,1.5], v in [0,2Pi]\n\n'
      '6. 圆锥面：x=u cos(v), y=u sin(v), z=u\n    u in [-2,2], v in [0,2Pi]')

add_h('四、程序设计', 13)
add_b('打开"数学实验程序.nb"，依次运行以下代码（光标放在每段代码上，按Shift+Enter运行）。')

add_b('曲面1 - 椭球面：')
add_code('ParametricPlot3D[{Sin[u]*Cos[v], Sin[u]*Sin[v], 2*Cos[u]},\n  {u, 0, Pi}, {v, 0, 2*Pi}, PlotPoints -> 30]')

add_b('曲面2 - 椭圆抛物面：')
add_code('ParametricPlot3D[{u*Cos[v], 2*u*Sin[v], u^2},\n  {u, 0, 2}, {v, 0, 2*Pi}, PlotPoints -> 30]')

add_b('曲面3 - 双曲抛物面（马鞍面）：')
add_code('Plot3D[x^2 - y^2, {x, -2, 2}, {y, -2, 2},\n  PlotPoints -> 30]')

add_b('曲面4 - 单叶双曲面：')
add_code('ParametricPlot3D[{Sec[u]*Cos[v], Sec[u]*Sin[v], Tan[u]},\n  {u, -Pi/4, Pi/4}, {v, 0, 2*Pi}, PlotPoints -> 30]')

add_b('曲面5 - 双叶双曲面：')
add_code('ParametricPlot3D[{Sqrt[1+u^2]*Cos[v], Sqrt[1+u^2]*Sin[v], u},\n  {u, -1.5, 1.5}, {v, 0, 2*Pi}, PlotPoints -> 30]')

add_b('曲面6 - 圆锥面：')
add_code('ParametricPlot3D[{u*Cos[v], u*Sin[v], u},\n  {u, -2, 2}, {v, 0, 2*Pi}, PlotPoints -> 30]')

add_h('五、程序运行结果', 13)
add_b('运行后得到六个三维曲面图形（用鼠标可以拖拽旋转，从不同角度观察）：')

add_pic('【图1-1 椭球面 - 请在此处插入运行截图】')
add_pic('【图1-2 椭圆抛物面 - 请在此处插入运行截图】')
add_pic('【图1-3 双曲抛物面（马鞍面）- 请在此处插入运行截图】')
add_pic('【图1-4 单叶双曲面 - 请在此处插入运行截图】')
add_pic('【图1-5 双叶双曲面 - 请在此处插入运行截图】')
add_pic('【图1-6 圆锥面 - 请在此处插入运行截图】')

add_h('六、结果的讨论和分析', 13)
add_b('通过这次实验，我对六种标准二次曲面有了直观的认识。以前在课堂上只能看课本上的截图，很难想象曲面的完整立体形状。在Mathematica中用鼠标拖拽旋转图形，从不同角度观察，对每种曲面的形状特征理解更深了。')

add_b('椭球面像一个压扁的橄榄球，z轴方向明显比水平方向长；椭圆抛物面像一个椭圆形的碗，开口朝上；马鞍面中间低、前后翘起，外形像马鞍，原点是一个鞍点——沿x轴方向向上弯，沿y轴方向向下弯；单叶双曲面中间细、两端粗，像一个腰鼓，是一个完整的连通曲面；双叶双曲面分为上下两片，像两个碗底对底放置；圆锥面是两个顶点对在一起的圆锥，关于xOy面对称。')

add_b('在操作上，ParametricPlot3D用于参数方程表示的曲面，Plot3D用于显函数z=f(x,y)形式的曲面。把参数方程填进去，设定好参数范围，Shift+Enter就能出图，操作很简便。')

doc.add_page_break()

# ===================================================================
# Experiment 2
# ===================================================================
add_h('实验二  最小二乘法（习题1）', 15)

add_h('一、实验题目', 13)
add_b('为测定刀具的磨损速度，每隔一小时测量一次刀具的厚度，得到以下数据。试根据这组数据建立y与t之间的拟合函数。')

add_b('时间t（小时）：0, 1, 2, 3, 4, 5, 6, 7\n'
      '厚度y（mm）：  27.0, 26.8, 26.5, 26.3, 26.1, 25.7, 25.3, 24.8')

add_h('二、实验目的和意义', 13)
add_b('在科学研究和实际工作中，常常需要根据实验数据找出变量之间的函数关系，即建立经验公式。本题中我需要根据刀具厚度随时间变化的数据，找出y与t之间的关系。最小二乘法是解决这类问题最常用的方法——找一条直线，使所有数据点到它的偏差平方和最小。')

add_h('三、计算公式', 13)
add_b('先画散点图观察趋势。数据点大致分布在一条直线附近，因此用线性函数拟合：y = a t + b\n\n'
      '最小二乘法原理：求a和b使偏差平方和 Q(a,b) = sum((a ti + b - yi)^2) 达到最小。\n\n'
      '由偏导数 dQ/da=0 和 dQ/db=0 得到正规方程组，解出a和b。\n\n'
      'Mathematica中可直接用Fit函数：Fit[data, {1, t}, t]')

add_h('四、程序设计', 13)
add_b('打开"数学实验程序.nb"，实验二部分，按顺序运行以下每段代码：')

add_b('步骤1 - 输入数据：')
add_code('t = {0, 1, 2, 3, 4, 5, 6, 7};\ny = {27.0, 26.8, 26.5, 26.3, 26.1, 25.7, 25.3, 24.8};\ndata = Table[{t[[i]], y[[i]]}, {i, 1, 8}]')

add_b('步骤2 - 画散点图，观察数据趋势：')
add_code('ListPlot[data, PlotStyle -> PointSize[0.02],\n  AxesLabel -> {"t (h)", "y (mm)"}]')

add_b('步骤3 - 用Fit函数做线性拟合（核心就这一行）：')
add_code('fit = Fit[data, {1, t}, t]')

add_b('步骤4 - 把散点图和拟合直线画在同一张图上：')
add_code('p1 = ListPlot[data, PlotStyle -> PointSize[0.02],\n    DisplayFunction -> Identity];\np2 = Plot[fit, {t, 0, 7}, DisplayFunction -> Identity,\n    PlotStyle -> RGBColor[1, 0, 0]];\nShow[p1, p2, DisplayFunction -> $DisplayFunction,\n  AxesLabel -> {"t (h)", "y (mm)"}]')

add_b('步骤5 - 循环计算每个时间点的拟合值和误差：')
add_code('Do[Print[t[[i]], "  ", y[[i]], "  ",\n  N[fit /. t -> t[[i]], 4], "  ",\n  N[y[[i]] - (fit /. t -> t[[i]]), 4]],\n  {i, 1, 8}]')

add_h('五、程序运行结果', 13)

add_pic('【图2-1 刀具磨损数据散点图 - 请在此处插入运行截图】')

add_b('散点图显示数据点大致沿一条下降直线分布，线性拟合是合适的。')

add_pic('【图2-2 Fit函数输出拟合结果 - 请在此处插入运行截图】')

add_b('得到拟合函数：y = 27.125 - 0.3036 t\n'
      '参数 a = -0.3036（斜率，刀具每小时磨损约0.30mm）\n'
      '参数 b = 27.125（截距，初始厚度约27.13mm）')

add_pic('【图2-3 散点图+拟合直线叠加图 - 请在此处插入运行截图】')

add_b('红色直线穿过数据点中间，点均匀分布在直线两侧，拟合良好。')

add_pic('【图2-4 拟合值与误差对比表 - 请在此处插入运行截图】')

add_b('拟合值与实际值对比：\n'
      't=0h: 实际27.0mm, 拟合27.13mm, 误差-0.13mm\n'
      't=1h: 实际26.8mm, 拟合26.82mm, 误差-0.02mm\n'
      't=2h: 实际26.5mm, 拟合26.52mm, 误差-0.02mm\n'
      't=3h: 实际26.3mm, 拟合26.21mm, 误差+0.09mm\n'
      't=4h: 实际26.1mm, 拟合25.91mm, 误差+0.19mm\n'
      't=5h: 实际25.7mm, 拟合25.61mm, 误差+0.09mm\n'
      't=6h: 实际25.3mm, 拟合25.30mm, 误差0.00mm\n'
      't=7h: 实际24.8mm, 拟合25.00mm, 误差-0.20mm\n\n'
      '最大误差约0.2mm，拟合效果良好。')

add_h('六、结果的讨论和分析', 13)
add_b('刀具磨损基本符合线性规律，每小时磨损约0.30mm。按此规律预测：8小时后约24.70mm，9小时后约24.39mm。')

add_b('体会：第一，拟合之前一定要先画散点图，判断适合用什么函数类型；第二，Fit函数使用方便，一行代码就完成最小二乘计算；第三，用Show把散点和拟合线叠加能直观检验拟合好坏；第四，输入数据时t和y要一一对应，顺序不能弄错。')

# Save
doc.save('report_final.docx')
print("Saved: report_final.docx")
