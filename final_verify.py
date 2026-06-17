#!/usr/bin/env python
# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn

docx_path = r'C:\Users\lenovo\Desktop\常用运算符和数学函数_含示例解释.docx'
doc = Document(docx_path)

print("=" * 60)
print("FINAL VERIFICATION - Formatting Comparison")
print("=" * 60)

# Compare original subsection (4.1) with new subsections (4.3, 4.4)
def describe_para(p, label):
    """Describe a paragraph's formatting."""
    r = p.runs[0] if p.runs else None
    if not r:
        return f"{label}: EMPTY"
    rPr = r._r.find(qn('w:rPr'))
    rFonts = rPr.find(qn('w:rFonts')) if rPr is not None else None
    color = rPr.find(qn('w:color')) if rPr is not None else None
    ascii_f = rFonts.get(qn('w:ascii')) if rFonts is not None else '-'
    ea_f = rFonts.get(qn('w:eastAsia')) if rFonts is not None else '-'
    c_val = color.get(qn('w:val')) if color is not None else '-'
    pPr = p._p.find(qn('w:pPr'))
    sp = pPr.find(qn('w:spacing')) if pPr is not None else None
    sp_str = 'no-spacing' if sp is None else f"b={sp.get(qn('w:before'))} a={sp.get(qn('w:after'))}"
    return (f"{label}: bold={r.bold}, size={r.font.size}, "
            f"font={ascii_f}/{ea_f}, color={c_val}, {sp_str}")

# 1. Subsection headers
print("\n--- Subsection Headers (4.1 original vs 4.3/4.4 new) ---")
for idx, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text in ['4.1 控制台 I/O（需 #include <iostream>）',
                 '4.3 I/O 操纵符（需 #include <iomanip>）',
                 '4.4 格式控制标志（ios::flags）']:
        print(f"  P{idx} {describe_para(p, text[:40])}")

# 2. Code paragraphs
print("\n--- Code Paragraphs (original vs new) ---")
# Original: P20 "// if-else"
p20 = doc.paragraphs[20]
print(f"  Original: {describe_para(p20, 'P20 code')}")
# New: first code line after 4.3
for idx, p in enumerate(doc.paragraphs):
    if p.text.strip().startswith('// 格式控制示例'):
        print(f"  New:      {describe_para(p, 'P' + str(idx) + ' code')}")
        break

# 3. Tables
print("\n--- Table Cells (original Table 0 vs new Table 9) ---")
def describe_table_cell(table_idx, row_idx, cell_idx, label):
    t = doc.tables[table_idx]
    cell = t.rows[row_idx].cells[cell_idx]
    runs_info = []
    for r in cell.paragraphs[0].runs:
        if r.text.strip():
            rPr = r._r.find(qn('w:rPr'))
            rFonts = rPr.find(qn('w:rFonts')) if rPr is not None else None
            ascii_f = rFonts.get(qn('w:ascii')) if rFonts is not None else '-'
            ea_f = rFonts.get(qn('w:eastAsia')) if rFonts is not None else '-'
            runs_info.append(f"'{r.text[:15]}' bold={r.bold} size={r.font.size} font={ascii_f}/{ea_f}")
    return f"{label}: {', '.join(runs_info) if runs_info else 'EMPTY'}"

print(f"  {describe_table_cell(0, 0, 0, 'Orig hdr')}")  # Original table header
print(f"  {describe_table_cell(9, 0, 0, 'New  hdr')}")  # New table header
print(f"  {describe_table_cell(0, 1, 0, 'Orig dat')}")  # Original table data
print(f"  {describe_table_cell(9, 1, 0, 'New  dat')}")  # New table data

# 4. Document structure summary
print("\n--- Document Structure ---")
print(f"Total paragraphs: {len(doc.paragraphs)}")
print(f"Total tables: {len(doc.tables)}")
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text and any(text.startswith(pfx) for pfx in ['一、','二、','三、','四、','五、','六、','七、','八、',
                                                        '4.1','4.2','4.3','4.4','4.5','4.6','4.7','4.8','4.9','4.10','4.11']):
        print(f"  P{i}: {text[:90]}")

print("\n" + "=" * 60)
print("VERIFICATION COMPLETE")
print("=" * 60)
