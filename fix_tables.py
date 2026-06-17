#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Fix table cells to remove extra empty runs."""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

docx_path = r'C:\Users\lenovo\Desktop\常用运算符和数学函数_含示例解释.docx'
doc = Document(docx_path)

FONT_NAME = '微软雅黑'
TABLE_SIZE = Pt(9)

def set_run_font(run, font_name=FONT_NAME, font_size=TABLE_SIZE, bold=False, color=None):
    run.font.name = font_name
    run.font.size = font_size
    run.bold = bold
    rPr = run._r.get_or_add_rPr()
    for old in rPr.findall(qn('w:rFonts')):
        rPr.remove(old)
    for old in rPr.findall(qn('w:color')):
        rPr.remove(old)
    rFonts = parse_xml(
        f'<w:rFonts {nsdecls("w")} '
        f'w:ascii="{font_name}" w:hAnsi="{font_name}" w:eastAsia="{font_name}"/>'
    )
    rPr.insert(0, rFonts)
    if color:
        rPr.append(parse_xml(f'<w:color {nsdecls("w")} w:val="{color}"/>'))

# Fix all new tables (tables 9-15 are the new ones)
for ti in range(9, 16):
    if ti >= len(doc.tables):
        break
    t = doc.tables[ti]
    print(f"Fixing Table {ti}...")

    # Fix header row
    for cell in t.rows[0].cells:
        for p in cell.paragraphs:
            # Remove empty runs
            for r in p.runs:
                if r.text.strip() == '':
                    r._r.getparent().remove(r._r)
                else:
                    set_run_font(r, font_size=TABLE_SIZE, bold=True)

    # Fix data rows
    for row in t.rows[1:]:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    if r.text.strip() == '':
                        r._r.getparent().remove(r._r)
                    else:
                        set_run_font(r, font_size=TABLE_SIZE, bold=False)

# Also remove extra empty runs from code paragraphs that might have them
for p in doc.paragraphs:
    empty_runs = [r for r in p.runs if r.text.strip() == '' and r.font.size is None]
    for r in empty_runs:
        try:
            r._r.getparent().remove(r._r)
        except:
            pass

doc.save(docx_path)
print("Fixed and saved.")
